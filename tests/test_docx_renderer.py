import base64
import subprocess
import sys
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Inches
from lxml import etree

import thesis_forge.renderers.docx.fields as fields_module
import thesis_forge.renderers.docx.renderer as renderer_module
import thesis_forge.renderers.docx.sections as sections_module
from thesis_forge.bibliography import Gbt7714Formatter, LocalBibTeXLoader
from thesis_forge.core.compiler import compile_document
from thesis_forge.core.math import MathSyntaxError
from thesis_forge.core.model import (
    Algorithm,
    BibliographyBlock,
    Citation,
    CrossReference,
    Equation,
    Figure,
    FootnoteDefinition,
    FootnoteReference,
    Heading,
    ListBlock,
    Listing,
    ListItem,
    Paragraph,
    Table,
    TableCell,
    TableRow,
    Text,
    ThesisDocument,
)
from thesis_forge.core.parser import parse_markdown
from thesis_forge.core.render_plan import (
    CitationRun,
    ParagraphInstruction,
    ReferenceRun,
    RenderPlan,
    TextRun,
    TocEntryInstruction,
    TocInstruction,
)
from thesis_forge.renderers.docx import DocxRenderer
from thesis_forge.renderers.docx.errors import DocxRenderError
from thesis_forge.renderers.docx.package import list_package_parts, read_package_part
from thesis_forge.renderers.docx.styles import (
    apply_paragraph_style,
    ensure_paragraph_style,
    resolve_paragraph_style,
)
from thesis_forge.templates import (
    AbstractStyleSpec,
    BibliographySpec,
    CoverSpec,
    DocumentGridSpec,
    FontSpec,
    LengthSpec,
    LineSpacingSpec,
    ListSpec,
    OrderedListLevelSpec,
    OrderedListSpec,
    ParagraphStyleSpec,
    SectionsSpec,
    TocLevelSpec,
    TocSpec,
    UnorderedListLevelSpec,
    UnorderedListSpec,
    load_template,
)

NS = {
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
}
REL_NS = {"pr": "http://schemas.openxmlformats.org/package/2006/relationships"}
PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk"
    "+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


def _text_inlines(value: str) -> list[Text]:
    return [Text(value=value)]


def _structured_table(
    table_id: str,
    caption: str,
    rows: list[tuple[bool, list[tuple[str, str | None]]]],
) -> Table:
    return Table(
        id=table_id,
        caption_inlines=tuple(_text_inlines(caption)),
        rows=tuple(
            TableRow(
                header=is_header,
                cells=tuple(
                    TableCell(
                        inlines=tuple(_text_inlines(value)),
                        alignment=alignment,
                    )
                    for value, alignment in cells
                ),
            )
            for is_header, cells in rows
        ),
    )


def _xml_part(path: Path, name: str):
    return etree.fromstring(read_package_part(path, name))


def _numbering_definition_for_paragraph(numbering_xml, paragraph):
    number_id = paragraph.xpath(
        "./w:pPr/w:numPr/w:numId/@w:val",
        namespaces=NS,
    )[0]
    abstract_id = numbering_xml.xpath(
        f"./w:num[@w:numId='{number_id}']/w:abstractNumId/@w:val",
        namespaces=NS,
    )[0]
    abstract = numbering_xml.xpath(
        f"./w:abstractNum[@w:abstractNumId='{abstract_id}']",
        namespaces=NS,
    )[0]
    return number_id, abstract


def test_docx_renderer_applies_template_page_body_and_heading_xml(tmp_path: Path):
    document = ThesisDocument(
        source_path=tmp_path / "thesis.md",
        blocks=[
            Heading(id="chap:intro", level=1, inlines=[Text(value="绪论")]),
            Paragraph(inlines=[Text(value="正文段落")]),
        ],
    )
    template = load_template("templates/schools/example-university/2026.yaml")
    plan = compile_document(document, template=template)
    output = tmp_path / "thesis.docx"

    DocxRenderer().render(plan, output)

    document_xml = _xml_part(output, "word/document.xml")
    styles_xml = _xml_part(output, "word/styles.xml")
    section = document_xml.find(".//w:sectPr", NS)
    assert section is not None
    page_size = section.find("w:pgSz", NS)
    margins = section.find("w:pgMar", NS)
    assert page_size is not None
    assert page_size.get(f"{{{NS['w']}}}w") == "11906"
    assert page_size.get(f"{{{NS['w']}}}h") == "16838"
    assert margins is not None
    assert margins.get(f"{{{NS['w']}}}left") == "1701"
    assert margins.get(f"{{{NS['w']}}}right") == "1417"

    normal = styles_xml.xpath(".//w:style[@w:styleId='Normal']", namespaces=NS)[0]
    assert normal.xpath("./w:rPr/w:rFonts/@w:eastAsia", namespaces=NS) == ["宋体"]
    assert normal.xpath("./w:rPr/w:rFonts/@w:ascii", namespaces=NS) == ["Times New Roman"]
    assert normal.xpath("./w:rPr/w:sz/@w:val", namespaces=NS) == ["24"]
    assert normal.xpath("./w:pPr/w:ind/@w:firstLine", namespaces=NS) == ["480"]
    assert normal.xpath("./w:pPr/w:spacing/@w:line", namespaces=NS) == ["400"]
    assert normal.xpath("./w:pPr/w:jc/@w:val", namespaces=NS) == ["both"]

    headings = [
        styles_xml.xpath(
            f".//w:style[@w:styleId='Heading{level}']",
            namespaces=NS,
        )[0]
        for level in range(1, 4)
    ]
    for heading in headings:
        assert heading.xpath("./w:rPr/w:rFonts/@w:eastAsia", namespaces=NS) == [
            "黑体"
        ]
        assert heading.xpath("./w:rPr/w:color/@w:val", namespaces=NS) == [
            "000000"
        ]
        assert not heading.xpath(
            "./w:rPr/w:color/@w:themeColor",
            namespaces=NS,
        )
        assert not heading.xpath(
            "./w:rPr/w:color/@w:themeTint",
            namespaces=NS,
        )
        assert not heading.xpath(
            "./w:rPr/w:color/@w:themeShade",
            namespaces=NS,
        )
        assert heading.xpath("./w:rPr/w:b", namespaces=NS)
    assert headings[0].xpath("./w:pPr/w:jc/@w:val", namespaces=NS) == ["center"]
    assert headings[0].xpath("./w:pPr/w:keepNext", namespaces=NS)
    assert headings[0].xpath("./w:pPr/w:keepLines", namespaces=NS)
    assert "[TODO:" not in etree.tostring(document_xml, encoding="unicode")
    assert "word/document.xml" in list_package_parts(output)


def test_docx_renderer_translates_complete_body_and_heading_policy_xml(
    tmp_path: Path,
):
    template = load_template("templates/base/bachelor.yaml")
    template.body.font = FontSpec(east_asia="楷体", latin="Arial")
    template.body.size = LengthSpec.model_validate("10pt")
    template.body.color = "auto"
    template.body.bold = True
    template.body.italic = True
    template.body.alignment = "right"
    template.body.left_indent = LengthSpec.model_validate("1em")
    template.body.right_indent = LengthSpec.model_validate("5pt")
    template.body.first_line_indent = LengthSpec.model_validate("2em")
    template.body.space_before = LengthSpec.model_validate("6pt")
    template.body.space_after = LengthSpec.model_validate("8pt")
    template.body.line_spacing = LineSpacingSpec(type="fixed", value="20pt")
    template.body.widow_control = False
    template.body.keep_together = True
    template.body.keep_with_next = False
    template.body.page_break_before = True
    template.body.outline_level = 2
    template.body.snap_to_grid = False

    heading = template.heading.level1
    heading.font = None
    heading.size = LengthSpec.model_validate("20pt")
    heading.color = "abcdef"
    heading.first_line_indent = None
    heading.hanging_indent = LengthSpec.model_validate("1.5em")
    heading.space_before = LengthSpec.model_validate("10pt")
    heading.space_after = LengthSpec.model_validate("4pt")
    heading.line_spacing = LineSpacingSpec(type="multiple", value=1.5)
    heading.widow_control = True
    heading.keep_together = False
    heading.keep_with_next = True
    heading.page_break_before = False
    heading.outline_level = 0
    heading.snap_to_grid = True

    document = ThesisDocument(
        source_path=tmp_path / "thesis.md",
        blocks=[
            Heading(
                id="chap:intro",
                level=1,
                inlines=_text_inlines("绪论"),
            ),
            Paragraph(inlines=[Text(value="正文")]),
        ],
    )
    output = tmp_path / "complete-policy.docx"

    DocxRenderer().render(compile_document(document, template=template), output)

    styles_xml = _xml_part(output, "word/styles.xml")
    normal = styles_xml.xpath(".//w:style[@w:styleId='Normal']", namespaces=NS)[0]
    assert normal.xpath("./w:rPr/w:rFonts/@w:eastAsia", namespaces=NS) == ["楷体"]
    assert normal.xpath("./w:rPr/w:rFonts/@w:ascii", namespaces=NS) == ["Arial"]
    assert normal.xpath("./w:rPr/w:sz/@w:val", namespaces=NS) == ["20"]
    assert normal.xpath("./w:rPr/w:color/@w:val", namespaces=NS) == ["auto"]
    assert normal.xpath("./w:rPr/w:b", namespaces=NS)
    assert normal.xpath("./w:rPr/w:i", namespaces=NS)
    assert normal.xpath("./w:pPr/w:jc/@w:val", namespaces=NS) == ["right"]
    assert normal.xpath("./w:pPr/w:ind/@w:left", namespaces=NS) == ["200"]
    assert normal.xpath("./w:pPr/w:ind/@w:right", namespaces=NS) == ["100"]
    assert normal.xpath("./w:pPr/w:ind/@w:firstLine", namespaces=NS) == ["400"]
    assert normal.xpath("./w:pPr/w:spacing/@w:before", namespaces=NS) == ["120"]
    assert normal.xpath("./w:pPr/w:spacing/@w:after", namespaces=NS) == ["160"]
    assert normal.xpath("./w:pPr/w:spacing/@w:line", namespaces=NS) == ["400"]
    assert normal.xpath("./w:pPr/w:spacing/@w:lineRule", namespaces=NS) == ["exact"]
    assert normal.xpath("./w:pPr/w:widowControl/@w:val", namespaces=NS) == ["0"]
    assert normal.xpath("./w:pPr/w:keepLines", namespaces=NS)
    assert normal.xpath("./w:pPr/w:keepNext/@w:val", namespaces=NS) == ["0"]
    assert normal.xpath("./w:pPr/w:pageBreakBefore", namespaces=NS)
    assert normal.xpath("./w:pPr/w:outlineLvl/@w:val", namespaces=NS) == ["2"]
    assert normal.xpath("./w:pPr/w:snapToGrid/@w:val", namespaces=NS) == ["0"]

    heading_xml = styles_xml.xpath(
        ".//w:style[@w:styleId='Heading1']",
        namespaces=NS,
    )[0]
    assert heading_xml.xpath("./w:rPr/w:rFonts/@w:eastAsia", namespaces=NS) == [
        "楷体"
    ]
    assert heading_xml.xpath("./w:rPr/w:rFonts/@w:ascii", namespaces=NS) == [
        "Arial"
    ]
    assert heading_xml.xpath("./w:rPr/w:sz/@w:val", namespaces=NS) == ["40"]
    assert heading_xml.xpath("./w:rPr/w:color/@w:val", namespaces=NS) == [
        "ABCDEF"
    ]
    assert not heading_xml.xpath("./w:rPr/w:color/@w:themeColor", namespaces=NS)
    assert not heading_xml.xpath("./w:rPr/w:color/@w:themeTint", namespaces=NS)
    assert not heading_xml.xpath("./w:rPr/w:color/@w:themeShade", namespaces=NS)
    assert heading_xml.xpath("./w:pPr/w:ind/@w:hanging", namespaces=NS) == ["600"]
    assert heading_xml.xpath("./w:pPr/w:spacing/@w:before", namespaces=NS) == ["200"]
    assert heading_xml.xpath("./w:pPr/w:spacing/@w:after", namespaces=NS) == ["80"]
    assert heading_xml.xpath("./w:pPr/w:spacing/@w:line", namespaces=NS) == ["360"]
    assert heading_xml.xpath("./w:pPr/w:spacing/@w:lineRule", namespaces=NS) == [
        "auto"
    ]
    assert heading_xml.xpath("./w:pPr/w:widowControl", namespaces=NS)
    assert heading_xml.xpath("./w:pPr/w:keepLines/@w:val", namespaces=NS) == ["0"]
    assert heading_xml.xpath("./w:pPr/w:keepNext", namespaces=NS)
    assert heading_xml.xpath("./w:pPr/w:pageBreakBefore/@w:val", namespaces=NS) == [
        "0"
    ]
    assert heading_xml.xpath("./w:pPr/w:outlineLvl/@w:val", namespaces=NS) == ["0"]
    assert heading_xml.xpath("./w:pPr/w:snapToGrid", namespaces=NS)


def test_paragraph_style_translator_uses_target_size_for_em_and_paragraph_runs():
    document = Document()
    paragraph = document.add_paragraph("正文")
    spec = ParagraphStyleSpec(
        font=FontSpec(east_asia="仿宋", latin="Calibri"),
        size="15pt",
        left_indent="2em",
        hanging_indent="1em",
        space_before="0.5em",
        line_spacing={"type": "fixed", "value": "2em"},
    )

    apply_paragraph_style(paragraph, spec)

    paragraph_xml = paragraph._p
    assert paragraph_xml.xpath("./w:pPr/w:ind/@w:left") == ["600"]
    assert paragraph_xml.xpath("./w:pPr/w:ind/@w:hanging") == ["300"]
    assert paragraph_xml.xpath("./w:pPr/w:spacing/@w:before") == ["150"]
    assert paragraph_xml.xpath("./w:pPr/w:spacing/@w:line") == ["600"]
    run_xml = paragraph.runs[0]._r
    assert run_xml.xpath("./w:rPr/w:rFonts/@w:eastAsia") == ["仿宋"]
    assert run_xml.xpath("./w:rPr/w:rFonts/@w:ascii") == ["Calibri"]
    assert run_xml.xpath("./w:rPr/w:sz/@w:val") == ["30"]


def test_heading_em_size_and_indent_resolve_from_body_font_size(tmp_path: Path):
    template = load_template("templates/base/bachelor.yaml")
    template.body.size = LengthSpec.model_validate("10pt")
    template.heading.level1.size = LengthSpec.model_validate("1.5em")
    template.heading.level1.left_indent = LengthSpec.model_validate("1em")
    document = ThesisDocument(
        source_path=tmp_path / "thesis.md",
        blocks=[
            Heading(
                id="chap:intro",
                level=1,
                inlines=_text_inlines("绪论"),
            )
        ],
    )
    output = tmp_path / "heading-em.docx"

    DocxRenderer().render(compile_document(document, template=template), output)

    styles_xml = _xml_part(output, "word/styles.xml")
    heading = styles_xml.xpath(
        ".//w:style[@w:styleId='Heading1']",
        namespaces=NS,
    )[0]
    assert heading.xpath("./w:rPr/w:sz/@w:val", namespaces=NS) == ["30"]
    assert heading.xpath("./w:pPr/w:ind/@w:left", namespaces=NS) == ["300"]


def test_single_line_spacing_writes_quantized_word_xml(tmp_path: Path):
    template = load_template("templates/base/bachelor.yaml")
    template.body.line_spacing = LineSpacingSpec(type="single")
    document = ThesisDocument(
        source_path=tmp_path / "thesis.md",
        blocks=[Paragraph(inlines=[Text(value="正文")])],
    )
    output = tmp_path / "single-spacing.docx"

    DocxRenderer().render(compile_document(document, template=template), output)

    styles_xml = _xml_part(output, "word/styles.xml")
    normal = styles_xml.xpath(".//w:style[@w:styleId='Normal']", namespaces=NS)[0]
    assert normal.xpath("./w:pPr/w:spacing/@w:lineRule", namespaces=NS) == ["auto"]
    assert normal.xpath("./w:pPr/w:spacing/@w:line", namespaces=NS) == ["240"]


@pytest.mark.parametrize(
    ("role", "style_name", "style_id"),
    [
        ("abstract.zh.title", "TF Abstract ZH Title", "TFAbstractZHTitle"),
        ("abstract.zh.body", "TF Abstract ZH Body", "TFAbstractZHBody"),
        ("keywords.zh", "TF Keywords ZH", "TFKeywordsZH"),
        ("abstract.en.title", "TF Abstract EN Title", "TFAbstractENTitle"),
        ("abstract.en.body", "TF Abstract EN Body", "TFAbstractENBody"),
        ("keywords.en", "TF Keywords EN", "TFKeywordsEN"),
        ("toc.title", "TF TOC Title", "TFTOCTitle"),
        (
            "bibliography.title",
            "TF Bibliography Title",
            "TFBibliographyTitle",
        ),
        (
            "bibliography.entry",
            "TF Bibliography Entry",
            "TFBibliographyEntry",
        ),
        (
            "special.acknowledgements",
            "TF Acknowledgements",
            "TFAcknowledgements",
        ),
        ("special.achievements", "TF Achievements", "TFAchievements"),
    ],
)
def test_ensure_paragraph_style_uses_stable_internal_style_ids(
    role: str,
    style_name: str,
    style_id: str,
):
    document = Document()
    spec = ParagraphStyleSpec(size="12pt", alignment="justify")

    first = ensure_paragraph_style(document, role, spec)
    second = ensure_paragraph_style(document, role, spec)

    assert first._element is second._element
    assert sum(
        style.style_id == style_id
        for style in document.styles
    ) == 1
    assert first.style_id == style_id
    assert first.name == style_name
    assert first.base_style.style_id == "Normal"


def test_ensure_paragraph_style_rejects_arbitrary_word_style_id():
    with pytest.raises(ValueError, match="unsupported paragraph role"):
        ensure_paragraph_style(
            Document(),
            "CustomWordStyle",
            ParagraphStyleSpec(size="12pt"),
        )


def test_stable_paragraph_style_survives_package_round_trip(tmp_path: Path):
    document = Document()
    style = ensure_paragraph_style(
        document,
        "abstract.zh.body",
        ParagraphStyleSpec(size="12pt", first_line_indent="2em"),
    )
    document.add_paragraph("摘要正文", style=style)
    output = tmp_path / "stable-style.docx"

    document.save(output)

    styles_xml = _xml_part(output, "word/styles.xml")
    document_xml = _xml_part(output, "word/document.xml")
    assert styles_xml.xpath(
        ".//w:style[@w:styleId='TFAbstractZHBody']",
        namespaces=NS,
    )
    assert document_xml.xpath(
        ".//w:p[.//w:t[text()='摘要正文']]/w:pPr/w:pStyle/@w:val",
        namespaces=NS,
    ) == ["TFAbstractZHBody"]

    reopened = Document(output)
    assert reopened.paragraphs[0].style.style_id == "TFAbstractZHBody"
    assert reopened.styles["TF Abstract ZH Body"].style_id == "TFAbstractZHBody"


def test_semantic_style_resolution_uses_deterministic_heading_and_body_fallbacks():
    template = load_template("templates/base/bachelor.yaml")

    assert resolve_paragraph_style(
        template,
        "abstract.zh.title",
        heading_level=1,
    ) is template.heading.level1
    assert resolve_paragraph_style(
        template,
        "abstract.zh.body",
    ) is template.body
    assert resolve_paragraph_style(
        template,
        "keywords.en",
    ) is template.body
    assert resolve_paragraph_style(
        template,
        "toc.title",
        heading_level=1,
    ) is template.heading.level1
    assert resolve_paragraph_style(
        template,
        "bibliography.entry",
    ) is template.body
    assert resolve_paragraph_style(
        template,
        "special.acknowledgements",
        heading_level=1,
    ) is template.heading.level1


def test_toc_level_styles_use_deterministic_defaults_and_real_field(tmp_path: Path):
    template = load_template("templates/base/bachelor.yaml")
    template.toc = TocSpec(title=ParagraphStyleSpec(size="16pt"))
    output = tmp_path / "default-toc-styles.docx"

    DocxRenderer().render(
        RenderPlan(nodes=[TocInstruction()], template=template),
        output,
    )

    styles_xml = _xml_part(output, "word/styles.xml")
    document_xml = _xml_part(output, "word/document.xml")
    settings_xml = _xml_part(output, "word/settings.xml")
    for level in range(1, 4):
        style = styles_xml.xpath(
            f".//w:style[@w:styleId='TOC{level}']",
            namespaces=NS,
        )[0]
        assert style.xpath("./w:name/@w:val", namespaces=NS) == [f"TOC {level}"]
        assert style.xpath("./w:basedOn/@w:val", namespaces=NS) == ["Normal"]
        assert style.xpath("./w:pPr/w:ind/@w:firstLine", namespaces=NS) == ["0"]
        assert style.xpath(
            "./w:pPr/w:tabs/w:tab/@w:val",
            namespaces=NS,
        ) == ["right"]
        assert style.xpath(
            "./w:pPr/w:tabs/w:tab/@w:pos",
            namespaces=NS,
        ) == ["8788"]
        assert style.xpath(
            "./w:pPr/w:tabs/w:tab/@w:leader",
            namespaces=NS,
        ) == ["dot"]

    field_codes = [
        "".join(node.itertext()).strip()
        for node in document_xml.xpath(".//w:instrText", namespaces=NS)
    ]
    assert field_codes == ['TOC \\o "1-3" \\h \\z \\u']
    toc_title = document_xml.xpath(
        ".//w:p[w:pPr/w:pStyle[@w:val='TFTOCTitle']]",
        namespaces=NS,
    )
    assert len(toc_title) == 1
    assert toc_title[0].xpath(".//w:t/text()", namespaces=NS) == ["目录"]
    assert toc_title[0].xpath(
        "./w:pPr/w:outlineLvl/@w:val",
        namespaces=NS,
    ) == ["9"]
    toc_field_paragraph = document_xml.xpath(
        ".//w:p[.//w:instrText[contains(., 'TOC')]]",
        namespaces=NS,
    )
    assert len(toc_field_paragraph) == 1
    assert toc_field_paragraph[0] is not toc_title[0]
    assert toc_field_paragraph[0].getprevious() is toc_title[0]
    assert toc_field_paragraph[0].xpath(".//w:t/text()", namespaces=NS) == []
    assert toc_field_paragraph[0].xpath(
        ".//w:bookmarkStart/@w:name",
        namespaces=NS,
    ) == ["tf_toc_index"]
    assert toc_field_paragraph[0].xpath(
        ".//w:bookmarkEnd/@w:id",
        namespaces=NS,
    ) == toc_field_paragraph[0].xpath(
        ".//w:bookmarkStart/@w:id",
        namespaces=NS,
    )
    assert document_xml.xpath(
        ".//w:instrText/../..//w:fldChar/@w:fldCharType",
        namespaces=NS,
    ) == ["begin", "separate", "end"]
    assert document_xml.xpath(
        ".//w:fldChar[@w:fldCharType='begin']/@w:dirty",
        namespaces=NS,
    ) == ["true"]
    assert settings_xml.xpath("./w:updateFields/@w:val", namespaces=NS) == ["true"]


def test_toc_level_styles_translate_indentation_spacing_tabs_and_leaders(
    tmp_path: Path,
):
    template = load_template("templates/base/bachelor.yaml")
    template.body.size = LengthSpec.model_validate("10pt")
    template.toc = TocSpec(
        title=ParagraphStyleSpec(size="16pt"),
        level1=TocLevelSpec(
            size="10pt",
            left_indent="0em",
            line_spacing={"type": "fixed", "value": "20pt"},
            page_number_tab="150mm",
            leader="dots",
        ),
        level2=TocLevelSpec(
            size="10pt",
            left_indent="1em",
            line_spacing={"type": "multiple", "value": 1.5},
            page_number_tab="145mm",
            leader="dashes",
        ),
        level3=TocLevelSpec(
            size="10pt",
            left_indent="2em",
            space_after="0.5em",
            page_number_tab="140mm",
            leader="middle_dot",
        ),
    )
    output = tmp_path / "configured-toc-styles.docx"

    DocxRenderer().render(
        RenderPlan(nodes=[TocInstruction()], template=template),
        output,
    )

    styles_xml = _xml_part(output, "word/styles.xml")
    level1 = styles_xml.xpath(
        ".//w:style[@w:styleId='TOC1']",
        namespaces=NS,
    )[0]
    level2 = styles_xml.xpath(
        ".//w:style[@w:styleId='TOC2']",
        namespaces=NS,
    )[0]
    level3 = styles_xml.xpath(
        ".//w:style[@w:styleId='TOC3']",
        namespaces=NS,
    )[0]

    assert level1.xpath("./w:pPr/w:ind/@w:left", namespaces=NS) == ["0"]
    assert level1.xpath("./w:pPr/w:spacing/@w:line", namespaces=NS) == ["400"]
    assert level1.xpath(
        "./w:pPr/w:spacing/@w:lineRule",
        namespaces=NS,
    ) == ["exact"]
    assert level2.xpath("./w:pPr/w:ind/@w:left", namespaces=NS) == ["200"]
    assert level2.xpath("./w:pPr/w:spacing/@w:line", namespaces=NS) == ["360"]
    assert level2.xpath(
        "./w:pPr/w:spacing/@w:lineRule",
        namespaces=NS,
    ) == ["auto"]
    assert level3.xpath("./w:pPr/w:ind/@w:left", namespaces=NS) == ["400"]
    assert level3.xpath("./w:pPr/w:spacing/@w:after", namespaces=NS) == ["100"]

    expected_tabs = {
        "TOC1": ("8504", "dot"),
        "TOC2": ("8220", "hyphen"),
        "TOC3": ("7937", "middleDot"),
    }
    for style_id, (position, leader) in expected_tabs.items():
        style = styles_xml.xpath(
            f".//w:style[@w:styleId={style_id!r}]",
            namespaces=NS,
        )[0]
        assert style.xpath(
            "./w:pPr/w:tabs/w:tab/@w:val",
            namespaces=NS,
        ) == ["right"]
        assert style.xpath(
            "./w:pPr/w:tabs/w:tab/@w:pos",
            namespaces=NS,
        ) == [position]
        assert style.xpath(
            "./w:pPr/w:tabs/w:tab/@w:leader",
            namespaces=NS,
        ) == [leader]


@pytest.mark.parametrize(
    ("leader", "word_value"),
    [
        ("none", "none"),
        ("dots", "dot"),
        ("dashes", "hyphen"),
        ("line", "underscore"),
        ("heavy", "heavy"),
        ("middle_dot", "middleDot"),
    ],
)
def test_toc_leader_policy_maps_to_word_tokens(
    tmp_path: Path,
    leader: str,
    word_value: str,
):
    template = load_template("templates/base/bachelor.yaml")
    template.toc = TocSpec(
        level1=TocLevelSpec(
            page_number_tab="150mm",
            leader=leader,
        )
    )
    output = tmp_path / f"toc-leader-{leader}.docx"

    DocxRenderer().render(
        RenderPlan(nodes=[TocInstruction()], template=template),
        output,
    )

    styles_xml = _xml_part(output, "word/styles.xml")
    assert styles_xml.xpath(
        ".//w:style[@w:styleId='TOC1']/w:pPr/w:tabs/w:tab/@w:leader",
        namespaces=NS,
    ) == [word_value]


def test_toc_em_page_number_tab_uses_effective_level_size(tmp_path: Path):
    template = load_template("templates/base/bachelor.yaml")
    template.body.size = LengthSpec.model_validate("12pt")
    template.toc = TocSpec(
        level1=TocLevelSpec(
            size="10pt",
            page_number_tab="10em",
        )
    )
    output = tmp_path / "toc-em-tab.docx"

    DocxRenderer().render(
        RenderPlan(nodes=[TocInstruction()], template=template),
        output,
    )

    styles_xml = _xml_part(output, "word/styles.xml")
    assert styles_xml.xpath(
        ".//w:style[@w:styleId='TOC1']/w:pPr/w:tabs/w:tab/@w:pos",
        namespaces=NS,
    ) == ["2000"]


def test_template_without_toc_policy_defines_toc_styles_with_deterministic_defaults(
    tmp_path: Path,
):
    # ADR-0005 §5.3：即使模板无 toc 配置也恒定义 TOC1-3（默认虚线右制表位），
    # 供 cached 条目与 LibreOffice 刷新引用，避免未定义样式引用。
    template = load_template("templates/base/bachelor.yaml")
    assert template.toc is None
    output = tmp_path / "legacy-toc-field.docx"

    DocxRenderer().render(
        RenderPlan(nodes=[TocInstruction()], template=template),
        output,
    )

    styles_xml = _xml_part(output, "word/styles.xml")
    document_xml = _xml_part(output, "word/document.xml")
    for level in range(1, 4):
        style = styles_xml.xpath(
            f".//w:style[@w:styleId='TOC{level}']",
            namespaces=NS,
        )[0]
        assert style.xpath("./w:name/@w:val", namespaces=NS) == [f"TOC {level}"]
        assert style.xpath("./w:basedOn/@w:val", namespaces=NS) == ["Normal"]
        assert style.xpath("./w:pPr/w:ind/@w:firstLine", namespaces=NS) == ["0"]
        assert style.xpath(
            "./w:pPr/w:tabs/w:tab/@w:val",
            namespaces=NS,
        ) == ["right"]
        assert style.xpath(
            "./w:pPr/w:tabs/w:tab/@w:pos",
            namespaces=NS,
        ) == ["8788"]
        assert style.xpath(
            "./w:pPr/w:tabs/w:tab/@w:leader",
            namespaces=NS,
        ) == ["dot"]
    assert [
        "".join(node.itertext()).strip()
        for node in document_xml.xpath(".//w:instrText", namespaces=NS)
    ] == ['TOC \\o "1-3" \\h \\z \\u']


def _toc_document(tmp_path: Path) -> ThesisDocument:
    return ThesisDocument(
        source_path=tmp_path / "thesis.md",
        blocks=[
            Heading(
                id="chap:abstract-zh",
                level=1,
                inlines=_text_inlines("摘要"),
            ),
            Heading(
                id="chap:intro",
                level=1,
                inlines=_text_inlines("绪论"),
            ),
            Heading(
                id="sec:background",
                level=2,
                inlines=_text_inlines("研究背景"),
            ),
            Heading(
                id="sec:limitations",
                level=3,
                inlines=_text_inlines("现有流程的局限"),
            ),
            Heading(
                id=None,
                level=2,
                inlines=_text_inlines("无书签小节"),
            ),
        ],
    )


def _toc_sections() -> SectionsSpec:
    return SectionsSpec.model_validate(
        {
            "front_matter": {
                "start": "new_page",
                "page_number": {"format": "roman-lower"},
            },
            "main": {
                "start": "new_page",
                "page_number": {"format": "decimal", "restart": 1},
            },
        }
    )


def test_compile_document_populates_toc_cached_entries(tmp_path: Path):
    template = load_template("templates/base/bachelor.yaml")
    template.sections = _toc_sections()

    plan = compile_document(_toc_document(tmp_path), template=template)

    toc = next(node for node in plan.nodes if isinstance(node, TocInstruction))
    assert toc.entries == (
        TocEntryInstruction("摘要", 1, "tf_chap_abstract_zh"),
        TocEntryInstruction("绪论", 1, "tf_chap_intro"),
        TocEntryInstruction("研究背景", 2, "tf_sec_background"),
        TocEntryInstruction("现有流程的局限", 3, "tf_sec_limitations"),
        TocEntryInstruction("无书签小节", 2, None),
    )
    assert toc.payload["entries"][0] == {
        "text": "摘要",
        "level": 1,
        "bookmark": "tf_chap_abstract_zh",
    }


def test_toc_cached_entries_render_real_word_cache_structure(tmp_path: Path):
    template = load_template("templates/base/bachelor.yaml")
    template.sections = _toc_sections()
    output = tmp_path / "toc-cached.docx"

    DocxRenderer().render(
        compile_document(_toc_document(tmp_path), template=template),
        output,
    )

    document_xml = _xml_part(output, "word/document.xml")
    field_paragraph = document_xml.xpath(
        ".//w:p[.//w:instrText[contains(., 'TOC')]]",
        namespaces=NS,
    )[0]
    # 字段段：bookmark + begin(dirty) + instr + separate，无正文文字
    assert field_paragraph.xpath(
        ".//w:bookmarkStart/@w:name",
        namespaces=NS,
    ) == ["tf_toc_index"]
    assert [
        "".join(node.itertext()).strip()
        for node in field_paragraph.xpath(".//w:instrText", namespaces=NS)
    ] == ['TOC \\o "1-3" \\h \\z \\u']
    assert field_paragraph.xpath(".//w:t/text()", namespaces=NS) == []
    assert field_paragraph.xpath(
        ".//w:fldChar[@w:fldCharType='begin']/@w:dirty",
        namespaces=NS,
    ) == ["true"]

    # cached 条目：字段段之后每条标题一段，样式按级别 TOC1/2/3
    body = field_paragraph.getparent()
    field_index = body.index(field_paragraph)
    entries = body[field_index + 1 : field_index + 6]
    assert [
        paragraph.xpath("./w:pPr/w:pStyle/@w:val", namespaces=NS)
        for paragraph in entries
    ] == [["TOC1"], ["TOC1"], ["TOC2"], ["TOC3"], ["TOC2"]]
    assert [
        "".join(paragraph.xpath(".//w:t/text()", namespaces=NS)[:1])
        for paragraph in entries
    ] == ["摘要", "绪论", "研究背景", "现有流程的局限", "无书签小节"]

    bookmark_names = set(
        document_xml.xpath(".//w:bookmarkStart/@w:name", namespaces=NS)
    )
    expected_anchors = {
        0: "tf_chap_abstract_zh",
        1: "tf_chap_intro",
        2: "tf_sec_background",
        3: "tf_sec_limitations",
    }
    for index, anchor in expected_anchors.items():
        hyperlink = entries[index].xpath("./w:hyperlink", namespaces=NS)[0]
        assert hyperlink.xpath("@w:anchor", namespaces=NS) == [anchor]
        assert anchor in bookmark_names
        # \h 行为：条目文字 + 右对齐制表位 + 内嵌 PAGEREF 字段（cached 占位 1）
        assert hyperlink.xpath("./w:r/w:tab", namespaces=NS)
        instructions = hyperlink.xpath(
            "./w:r/w:instrText/text()",
            namespaces=NS,
        )
        assert instructions == [f"PAGEREF {anchor} \\h"]
        field_chars = hyperlink.xpath(
            "./w:r/w:fldChar/@w:fldCharType",
            namespaces=NS,
        )
        assert field_chars == ["begin", "separate", "end"]
        assert hyperlink.xpath(
            "./w:r/w:fldChar[@w:fldCharType='begin']/@w:dirty",
            namespaces=NS,
        ) == ["true"]
        pageref_cached = "".join(hyperlink.xpath(".//w:t/text()", namespaces=NS))
        assert pageref_cached.endswith("1")

    # 无书签标题降级为纯文本占位（无 hyperlink / PAGEREF）
    anonymous = entries[4]
    assert not anonymous.xpath("./w:hyperlink", namespaces=NS)
    assert not anonymous.xpath(".//w:instrText", namespaces=NS)
    assert anonymous.xpath("./w:r/w:tab", namespaces=NS)
    assert anonymous.xpath(".//w:t/text()", namespaces=NS) == ["无书签小节", "1"]

    # TOC 字段 end 落在最后一个条目段末尾（Word 原生结构）
    last = entries[-1]
    assert last.xpath("./w:r/w:fldChar/@w:fldCharType", namespaces=NS) == ["end"]
    assert not entries[-2].xpath(
        "./w:r/w:fldChar[@w:fldCharType='end']",
        namespaces=NS,
    )

    # 全部字段（含嵌套 PAGEREF）配对平衡
    field_types = document_xml.xpath(".//w:fldChar/@w:fldCharType", namespaces=NS)
    assert field_types.count("begin") == field_types.count("separate")
    assert field_types.count("begin") == field_types.count("end")


def test_toc_cached_entries_pass_openxml_validate(tmp_path: Path):
    template = load_template("templates/base/bachelor.yaml")
    template.sections = _toc_sections()
    output = tmp_path / "toc-cached-validate.docx"
    DocxRenderer().render(
        compile_document(_toc_document(tmp_path), template=template),
        output,
    )

    result = subprocess.run(
        [
            sys.executable,
            str(Path(__file__).resolve().parents[1] / "qa" / "tools" / "openxml_validate.py"),
            str(output),
        ],
        capture_output=True,
        text=True,
        check=False,
    )
    assert result.returncode == 0, result.stdout


def test_partial_semantic_title_inherits_heading_style_and_overrides_false(
    tmp_path: Path,
):
    template = load_template("templates/base/bachelor.yaml")
    template.heading.level1.space_before = LengthSpec.model_validate("10pt")
    template.heading.level1.keep_with_next = True
    template.heading.level1.page_break_before = True
    template.semantic_styles.abstract_zh = AbstractStyleSpec(
        title=ParagraphStyleSpec(
            space_before="1em",
            page_break_before=False,
        )
    )
    document = ThesisDocument(
        source_path=tmp_path / "thesis.md",
        blocks=[
            Heading(
                id="chap:abstract-zh",
                level=1,
                inlines=_text_inlines("摘要"),
            )
        ],
    )
    output = tmp_path / "partial-semantic-title.docx"

    DocxRenderer().render(compile_document(document, template=template), output)

    styles_xml = _xml_part(output, "word/styles.xml")
    heading = styles_xml.xpath(
        ".//w:style[@w:styleId='Heading1']",
        namespaces=NS,
    )[0]
    semantic = styles_xml.xpath(
        ".//w:style[@w:styleId='TFAbstractZHTitle']",
        namespaces=NS,
    )[0]
    assert heading.xpath("./w:rPr/w:b", namespaces=NS)
    assert heading.xpath("./w:pPr/w:jc/@w:val", namespaces=NS) == ["center"]
    assert heading.xpath("./w:pPr/w:spacing/@w:before", namespaces=NS) == ["200"]
    assert heading.xpath("./w:pPr/w:keepNext", namespaces=NS)
    assert heading.xpath("./w:pPr/w:pageBreakBefore", namespaces=NS)
    assert semantic.xpath("./w:basedOn/@w:val", namespaces=NS) == ["Heading1"]
    assert not semantic.xpath("./w:rPr/w:b", namespaces=NS)
    assert not semantic.xpath("./w:rPr/w:rFonts", namespaces=NS)
    assert not semantic.xpath("./w:rPr/w:sz", namespaces=NS)
    assert not semantic.xpath("./w:pPr/w:jc", namespaces=NS)
    assert semantic.xpath(
        "./w:pPr/w:spacing/@w:before",
        namespaces=NS,
    ) == ["320"]
    assert not semantic.xpath("./w:pPr/w:keepNext", namespaces=NS)
    assert semantic.xpath(
        "./w:pPr/w:pageBreakBefore/@w:val",
        namespaces=NS,
    ) == ["0"]


def test_partial_semantic_body_uses_inherited_size_for_em_lengths(
    tmp_path: Path,
):
    template = load_template("templates/base/bachelor.yaml")
    template.semantic_styles.abstract_zh = AbstractStyleSpec(
        body=ParagraphStyleSpec(
            first_line_indent="2em",
            space_after="0.5em",
            line_spacing={"type": "fixed", "value": "1.5em"},
        )
    )
    document = ThesisDocument(
        source_path=tmp_path / "thesis.md",
        blocks=[
            Heading(
                id="chap:abstract-zh",
                level=1,
                inlines=_text_inlines("摘要"),
            ),
            Paragraph(inlines=_text_inlines("摘要正文")),
        ],
    )
    output = tmp_path / "partial-semantic-body.docx"

    DocxRenderer().render(compile_document(document, template=template), output)

    styles_xml = _xml_part(output, "word/styles.xml")
    semantic = styles_xml.xpath(
        ".//w:style[@w:styleId='TFAbstractZHBody']",
        namespaces=NS,
    )[0]
    assert semantic.xpath("./w:basedOn/@w:val", namespaces=NS) == ["Normal"]
    assert not semantic.xpath("./w:rPr/w:rFonts", namespaces=NS)
    assert not semantic.xpath("./w:rPr/w:sz", namespaces=NS)
    assert semantic.xpath("./w:pPr/w:ind/@w:firstLine", namespaces=NS) == ["480"]
    assert semantic.xpath("./w:pPr/w:spacing/@w:after", namespaces=NS) == ["120"]
    assert semantic.xpath("./w:pPr/w:spacing/@w:line", namespaces=NS) == ["360"]
    assert semantic.xpath(
        "./w:pPr/w:spacing/@w:lineRule",
        namespaces=NS,
    ) == ["exact"]


def test_docx_renderer_binds_complete_abstract_fragment_to_semantic_styles(
    tmp_path: Path,
):
    source = tmp_path / "semantic-fragment.md"
    source.write_text(
        """# 摘要 {#chap:abstract-zh}

中文摘要正文。

关键词：编译；模板

# Abstract {#chap:abstract-en}

English abstract body.

Keywords: compiler; template

# 目录 {#chap:toc}

# 参考文献 {#references}

[1] Reference entry.

# 致谢 {#acknowledgements}

感谢所有帮助。

# 攻读学位期间的成果 {#achievements}
""",
        encoding="utf-8",
    )
    template = load_template("templates/base/bachelor.yaml")
    template.semantic_styles.abstract_zh = AbstractStyleSpec(
        title=ParagraphStyleSpec(size="18pt", alignment="center"),
        body=ParagraphStyleSpec(size="12pt", first_line_indent="2em"),
        keywords=ParagraphStyleSpec(size="11pt", first_line_indent="0em"),
    )
    template.semantic_styles.abstract_en = AbstractStyleSpec(
        title=ParagraphStyleSpec(size="17pt", alignment="center"),
        body=ParagraphStyleSpec(size="10pt", first_line_indent="1em"),
        keywords=ParagraphStyleSpec(size="9pt", first_line_indent="0em"),
    )
    template.bibliography = BibliographySpec(
        title=ParagraphStyleSpec(size="16pt", alignment="center"),
        entry=ParagraphStyleSpec(size="10pt", hanging_indent="2em"),
    )
    template.semantic_styles.acknowledgements = ParagraphStyleSpec(
        size="16pt",
        alignment="center",
    )
    template.semantic_styles.achievements = ParagraphStyleSpec(
        size="15pt",
        alignment="center",
    )
    template.toc = TocSpec(title=ParagraphStyleSpec(size="16pt"))
    output = tmp_path / "semantic-fragment.docx"

    DocxRenderer().render(
        compile_document(parse_markdown(source), template=template),
        output,
    )

    document_xml = _xml_part(output, "word/document.xml")
    styles_xml = _xml_part(output, "word/styles.xml")
    expected_styles = {
        "摘要": "TFAbstractZHTitle",
        "中文摘要正文。": "TFAbstractZHBody",
        "关键词：编译；模板": "TFKeywordsZH",
        "Abstract": "TFAbstractENTitle",
        "English abstract body.": "TFAbstractENBody",
        "Keywords: compiler; template": "TFKeywordsEN",
        "目录": "TFTOCTitle",
        "参考文献": "TFBibliographyTitle",
        "[1] Reference entry.": "TFBibliographyEntry",
        "致谢": "TFAcknowledgements",
        "攻读学位期间的成果": "TFAchievements",
    }
    for text, style_id in expected_styles.items():
        assert document_xml.xpath(
            f".//w:p[.//w:t[text()={text!r}]]/w:pPr/w:pStyle/@w:val",
            namespaces=NS,
        ) == [style_id]
        assert styles_xml.xpath(
            f".//w:style[@w:styleId={style_id!r}]",
            namespaces=NS,
        )

    assert styles_xml.xpath(
        ".//w:style[@w:styleId='TFAbstractZHBody']/w:pPr/w:ind/@w:firstLine",
        namespaces=NS,
    ) == ["480"]
    assert styles_xml.xpath(
        ".//w:style[@w:styleId='TFKeywordsEN']/w:rPr/w:sz/@w:val",
        namespaces=NS,
    ) == ["18"]


def test_heading_levels_one_through_three_use_shared_translator(tmp_path: Path):
    template = load_template("templates/base/bachelor.yaml")
    expected_left_indents = {"Heading1": "320", "Heading2": "560", "Heading3": "720"}
    for level in range(1, 4):
        heading = template.heading.for_level(level)
        assert heading is not None
        heading.left_indent = LengthSpec.model_validate(f"{level}em")
        heading.outline_level = level - 1
        heading.keep_with_next = True
        heading.snap_to_grid = False
    document = ThesisDocument(
        source_path=tmp_path / "thesis.md",
        blocks=[
            Heading(
                id="chap:one",
                level=1,
                inlines=_text_inlines("一级标题"),
            ),
            Heading(
                id="sec:two",
                level=2,
                inlines=_text_inlines("二级标题"),
            ),
            Heading(
                id="sec:three",
                level=3,
                inlines=_text_inlines("三级标题"),
            ),
        ],
    )
    output = tmp_path / "heading-levels.docx"

    DocxRenderer().render(compile_document(document, template=template), output)

    styles_xml = _xml_part(output, "word/styles.xml")
    for level in range(1, 4):
        style_id = f"Heading{level}"
        style = styles_xml.xpath(
            f".//w:style[@w:styleId='{style_id}']",
            namespaces=NS,
        )[0]
        assert style.xpath("./w:pPr/w:ind/@w:left", namespaces=NS) == [
            expected_left_indents[style_id]
        ]
        assert style.xpath("./w:pPr/w:outlineLvl/@w:val", namespaces=NS) == [
            str(level - 1)
        ]
        assert style.xpath("./w:pPr/w:keepNext", namespaces=NS)
        assert style.xpath("./w:pPr/w:snapToGrid/@w:val", namespaces=NS) == ["0"]


def test_two_templates_change_styles_without_changing_document_semantics(
    tmp_path: Path,
):
    document = ThesisDocument(
        source_path=tmp_path / "thesis.md",
        blocks=[
            Heading(
                id="chap:intro",
                level=1,
                inlines=_text_inlines("绪论"),
            ),
            Paragraph(inlines=[Text(value="相同正文")]),
        ],
    )
    first_template = load_template("templates/base/bachelor.yaml")
    second_template = load_template("templates/base/bachelor.yaml")
    second_template.body.font = FontSpec(east_asia="楷体", latin="Arial")
    second_template.body.size = LengthSpec.model_validate("10pt")
    second_template.body.first_line_indent = LengthSpec.model_validate("3em")
    first_output = tmp_path / "first.docx"
    second_output = tmp_path / "second.docx"

    DocxRenderer().render(
        compile_document(document, template=first_template),
        first_output,
    )
    DocxRenderer().render(
        compile_document(document, template=second_template),
        second_output,
    )

    first_document = _xml_part(first_output, "word/document.xml")
    second_document = _xml_part(second_output, "word/document.xml")
    assert first_document.xpath(".//w:body//w:t/text()", namespaces=NS) == (
        second_document.xpath(".//w:body//w:t/text()", namespaces=NS)
    )
    first_styles = _xml_part(first_output, "word/styles.xml")
    second_styles = _xml_part(second_output, "word/styles.xml")
    first_normal = first_styles.xpath(
        ".//w:style[@w:styleId='Normal']",
        namespaces=NS,
    )[0]
    second_normal = second_styles.xpath(
        ".//w:style[@w:styleId='Normal']",
        namespaces=NS,
    )[0]
    assert etree.tostring(first_normal) != etree.tostring(second_normal)
    assert second_normal.xpath("./w:rPr/w:rFonts/@w:eastAsia", namespaces=NS) == [
        "楷体"
    ]
    assert second_normal.xpath("./w:pPr/w:ind/@w:firstLine", namespaces=NS) == [
        "600"
    ]


def test_docx_renderer_applies_landscape_orientation(tmp_path: Path):
    template = load_template("templates/base/bachelor.yaml")
    template.page.orientation = "landscape"
    document = ThesisDocument(
        source_path=tmp_path / "thesis.md",
        blocks=[Paragraph(inlines=[Text(value="正文")])],
    )
    output = tmp_path / "landscape.docx"

    DocxRenderer().render(compile_document(document, template=template), output)

    document_xml = _xml_part(output, "word/document.xml")
    page_size = document_xml.find(".//w:sectPr/w:pgSz", NS)
    assert page_size is not None
    assert page_size.get(f"{{{NS['w']}}}orient") == "landscape"
    assert page_size.get(f"{{{NS['w']}}}w") == "16838"
    assert page_size.get(f"{{{NS['w']}}}h") == "11906"


def test_docx_renderer_writes_metadata_cover_before_front_matter(tmp_path: Path):
    template = load_template("templates/schools/example-university/2026.yaml")
    document = ThesisDocument(
        source_path=tmp_path / "thesis.md",
        metadata={
            "university": {"name": "XX大学", "college": "计算机学院"},
            "thesis": {
                "title": "结构化论文编译",
                "title_en": "Structured Thesis Compilation",
                "major": "计算机科学与技术",
                "degree": "工学学士",
            },
            "author": {"name": "张三", "student_id": "2022000001"},
            "advisor": {"name": "李老师", "title": "副教授"},
            "dates": {"completed": "2026-06"},
        },
        blocks=[
            Heading(
                id="chap:abstract-zh",
                level=1,
                inlines=_text_inlines("摘要"),
            ),
            Heading(
                id="chap:introduction",
                level=1,
                inlines=_text_inlines("绪论"),
            ),
        ],
    )
    output = tmp_path / "cover.docx"

    DocxRenderer().render(compile_document(document, template=template), output)

    document_xml = _xml_part(output, "word/document.xml")
    body_text = "".join(document_xml.xpath(".//w:body//w:t/text()", namespaces=NS))
    assert body_text.startswith(
        "XX大学计算机学院结构化论文编译Structured Thesis Compilation"
    )
    for value in (
        "计算机科学与技术",
        "工学学士",
        "张三",
        "2022000001",
        "李老师",
        "副教授",
        "2026-06",
        "摘要",
        "目录",
        "绪论",
    ):
        assert value in body_text
    for cover_text in ("XX大学", "结构化论文编译"):
        cover_paragraph = document_xml.xpath(
            f".//w:p[.//w:t[text()='{cover_text}']]",
            namespaces=NS,
        )[0]
        assert not any(
            style.startswith("Heading")
            for style in cover_paragraph.xpath(
                "./w:pPr/w:pStyle/@w:val",
                namespaces=NS,
            )
        )
    # 标题文本同时出现在 TOC cached 条目（TOC1 样式）中，需按样式区分
    assert len(
        document_xml.xpath(
            ".//w:p[w:pPr/w:pStyle[@w:val='TFAbstractZHTitle']]"
            "[.//w:t[text()='摘要']]",
            namespaces=NS,
        )
    ) == 1
    assert len(document_xml.xpath(".//w:sectPr", namespaces=NS)) == 3
    assert document_xml.xpath(".//w:headerReference", namespaces=NS)
    assert document_xml.xpath(".//w:footerReference", namespaces=NS)


def test_docx_renderer_uses_template_cover_order_content_and_style(tmp_path: Path):
    template = load_template("templates/schools/example-university/2026.yaml")
    template.cover = CoverSpec.model_validate(
        {
            "items": [
                {
                    "field": "thesis.title",
                    "prefix": "题目：",
                    "style": {
                        "font": {
                            "east_asia": "黑体",
                            "latin": "Arial",
                        },
                        "size": "18pt",
                        "color": "123456",
                        "bold": True,
                        "alignment": "right",
                        "space_before": "10pt",
                        "space_after": "12pt",
                    },
                },
                {
                    "text": "硕士学位论文",
                    "style": {
                        "alignment": "center",
                    },
                },
                {
                    "field": "advisor.title",
                    "prefix": "导师职称：",
                    "skip_if_empty": True,
                },
            ]
        }
    )
    document = ThesisDocument(
        source_path=tmp_path / "thesis.md",
        metadata={
            "thesis": {"title": "参数化封面"},
        },
        blocks=[
            Heading(
                id="chap:introduction",
                level=1,
                inlines=_text_inlines("绪论"),
            )
        ],
    )
    output = tmp_path / "parameterized-cover.docx"

    DocxRenderer().render(compile_document(document, template=template), output)

    document_xml = _xml_part(output, "word/document.xml")
    paragraphs = document_xml.xpath(".//w:body/w:p", namespaces=NS)
    paragraph_text = [
        "".join(paragraph.xpath(".//w:t/text()", namespaces=NS))
        for paragraph in paragraphs
    ]
    assert paragraph_text[:2] == ["题目：参数化封面", "硕士学位论文"]
    assert "导师职称：" not in paragraph_text

    title = paragraphs[0]
    assert title.xpath("./w:pPr/w:jc/@w:val", namespaces=NS) == ["right"]
    assert title.xpath("./w:pPr/w:spacing/@w:before", namespaces=NS) == ["200"]
    assert title.xpath("./w:pPr/w:spacing/@w:after", namespaces=NS) == ["240"]
    assert title.xpath("./w:r/w:rPr/w:rFonts/@w:eastAsia", namespaces=NS) == [
        "黑体"
    ]
    assert title.xpath("./w:r/w:rPr/w:rFonts/@w:ascii", namespaces=NS) == [
        "Arial"
    ]
    assert title.xpath("./w:r/w:rPr/w:sz/@w:val", namespaces=NS) == ["36"]
    assert title.xpath("./w:r/w:rPr/w:color/@w:val", namespaces=NS) == [
        "123456"
    ]
    assert title.xpath("./w:r/w:rPr/w:b", namespaces=NS)


def test_docx_renderer_bookmarks_listing_and_algorithm_objects(tmp_path: Path):
    document = ThesisDocument(
        source_path=tmp_path / "thesis.md",
        blocks=[
            Listing(
                id="lst:service",
                caption_inlines=_text_inlines("构建服务"),
                language="python",
                code="build_service(source, output)",
            ),
            Algorithm(
                id="alg:build",
                caption_inlines=_text_inlines("安全构建"),
                body="1. validate\n2. render\n3. replace",
            ),
        ],
    )
    output = tmp_path / "listing-algorithm.docx"

    DocxRenderer().render(
        compile_document(
            document,
            template=load_template("templates/base/bachelor.yaml"),
        ),
        output,
    )

    document_xml = _xml_part(output, "word/document.xml")
    bookmark_starts = {
        node.get(f"{{{NS['w']}}}id"): node.get(f"{{{NS['w']}}}name")
        for node in document_xml.xpath(".//w:bookmarkStart", namespaces=NS)
    }
    bookmark_ends = set(
        document_xml.xpath(".//w:bookmarkEnd/@w:id", namespaces=NS)
    )
    assert {"tf_lst_service", "tf_alg_build"} <= set(bookmark_starts.values())
    assert set(bookmark_starts) <= bookmark_ends
    assert document_xml.xpath(
        ".//w:p[.//w:bookmarkStart[@w:name='tf_lst_service']]//w:t[text()='构建服务']",
        namespaces=NS,
    )
    assert document_xml.xpath(
        ".//w:p[.//w:bookmarkStart[@w:name='tf_alg_build']]//w:t[text()='安全构建']",
        namespaces=NS,
    )
    for body_text in ("build_service(source, output)", "1. validate"):
        body_paragraph = document_xml.xpath(
            f".//w:p[.//w:t[contains(., '{body_text}')]]",
            namespaces=NS,
        )[0]
        assert body_paragraph.xpath("./w:pPr/w:jc/@w:val", namespaces=NS) == ["left"]
        assert body_paragraph.xpath(
            "./w:pPr/w:spacing/@w:lineRule",
            namespaces=NS,
        ) == ["auto"]
        assert body_paragraph.xpath(
            "./w:pPr/w:spacing/@w:line",
            namespaces=NS,
        ) == ["240"]


def test_docx_renderer_preserves_list_start_and_nesting_xml(tmp_path: Path):
    document = ThesisDocument(
        source_path=tmp_path / "thesis.md",
        blocks=[
            ListBlock(
                ordered=True,
                start=3,
                items=[
                    ListItem(level=0, ordinal=3, inlines=[Text(value="第三项")]),
                    ListItem(level=1, ordinal=1, inlines=[Text(value="子项")]),
                    ListItem(level=0, ordinal=4, inlines=[Text(value="第四项")]),
                ],
            )
        ],
    )
    output = tmp_path / "list.docx"

    DocxRenderer().render(
        compile_document(document, template=load_template("templates/base/bachelor.yaml")),
        output,
    )

    document_xml = _xml_part(output, "word/document.xml")
    numbering_xml = _xml_part(output, "word/numbering.xml")
    list_paragraphs = document_xml.xpath(
        ".//w:p[w:pPr/w:numPr]",
        namespaces=NS,
    )
    assert [
        paragraph.xpath("./w:pPr/w:numPr/w:ilvl/@w:val", namespaces=NS)[0]
        for paragraph in list_paragraphs
    ] == [
        "0",
        "1",
        "0",
    ]
    number_ids = {
        paragraph.xpath(
            "./w:pPr/w:numPr/w:numId/@w:val",
            namespaces=NS,
        )[0]
        for paragraph in list_paragraphs
    }
    assert len(number_ids) == 1
    _, abstract = _numbering_definition_for_paragraph(
        numbering_xml,
        list_paragraphs[0],
    )
    assert abstract.xpath(
        "./w:lvl[@w:ilvl='0']/w:start/@w:val",
        namespaces=NS,
    ) == ["3"]
    assert abstract.xpath(
        "./w:lvl/w:numFmt/@w:val",
        namespaces=NS,
    ) == ["decimal"] * 9
    assert abstract.xpath(
        "./w:lvl/w:lvlText/@w:val",
        namespaces=NS,
    ) == [f"%{level}." for level in range(1, 10)]
    children = [etree.QName(child).localname for child in numbering_xml]
    assert max(index for index, name in enumerate(children) if name == "abstractNum") < min(
        index for index, name in enumerate(children) if name == "num"
    )
    abstract_ids = set(
        numbering_xml.xpath("./w:abstractNum/@w:abstractNumId", namespaces=NS)
    )
    referenced_ids = set(
        numbering_xml.xpath("./w:num/w:abstractNumId/@w:val", namespaces=NS)
    )
    assert referenced_ids <= abstract_ids


def test_docx_renderer_applies_independent_template_list_policies_and_styles(
    tmp_path: Path,
):
    template = load_template("templates/base/bachelor.yaml")
    list_style = ParagraphStyleSpec(
        font=FontSpec(east_asia="楷体", latin="Arial"),
        size=LengthSpec.model_validate("11pt"),
        color="123456",
        bold=True,
        italic=True,
        alignment="justify",
        space_before=LengthSpec.model_validate("6pt"),
        space_after=LengthSpec.model_validate("8pt"),
        line_spacing=LineSpacingSpec(type="fixed", value="18pt"),
    )
    template.list = ListSpec(
        ordered=OrderedListSpec(
            levels=(
                OrderedListLevelSpec(
                    format="lower_roman",
                    prefix="(",
                    suffix=")",
                    alignment="right",
                    left_indent="42pt",
                    hanging_indent="12pt",
                    style=list_style,
                ),
                OrderedListLevelSpec(
                    format="upper_letter",
                    prefix="[",
                    suffix="]",
                    alignment="center",
                    left_indent="60pt",
                    hanging_indent="18pt",
                    style=list_style,
                ),
            )
        ),
        unordered=UnorderedListSpec(
            levels=(
                UnorderedListLevelSpec(
                    marker="◆",
                    alignment="center",
                    left_indent="30pt",
                    hanging_indent="10pt",
                    style=ParagraphStyleSpec(
                        font=FontSpec(east_asia="仿宋", latin="Courier New"),
                        size=LengthSpec.model_validate("10pt"),
                        color="654321",
                        space_after=LengthSpec.model_validate("4pt"),
                    ),
                ),
            )
        ),
    )
    document = ThesisDocument(
        source_path=tmp_path / "thesis.md",
        blocks=[
            ListBlock(
                ordered=True,
                start=3,
                items=[
                    ListItem(
                        level=0,
                        ordinal=3,
                        inlines=[Text(value="第三项")],
                    ),
                    ListItem(
                        level=12,
                        ordinal=1,
                        inlines=[Text(value="深层项")],
                    ),
                ],
            ),
            ListBlock(
                ordered=False,
                items=[
                    ListItem(
                        level=0,
                        inlines=[Text(value="项目符号")],
                    ),
                ],
            ),
        ],
    )
    output = tmp_path / "custom-list.docx"

    DocxRenderer().render(compile_document(document, template=template), output)

    document_xml = _xml_part(output, "word/document.xml")
    numbering_xml = _xml_part(output, "word/numbering.xml")
    paragraphs = document_xml.xpath(".//w:p[w:pPr/w:numPr]", namespaces=NS)
    assert len(paragraphs) == 3
    ordered_num_id, ordered_abstract = _numbering_definition_for_paragraph(
        numbering_xml,
        paragraphs[0],
    )
    deep_num_id, deep_abstract = _numbering_definition_for_paragraph(
        numbering_xml,
        paragraphs[1],
    )
    unordered_num_id, unordered_abstract = _numbering_definition_for_paragraph(
        numbering_xml,
        paragraphs[2],
    )
    assert ordered_num_id == deep_num_id
    assert ordered_abstract is deep_abstract
    assert unordered_num_id != ordered_num_id
    assert paragraphs[1].xpath(
        "./w:pPr/w:numPr/w:ilvl/@w:val",
        namespaces=NS,
    ) == ["8"]

    ordered_level = ordered_abstract.xpath(
        "./w:lvl[@w:ilvl='0']",
        namespaces=NS,
    )[0]
    assert ordered_level.xpath("./w:start/@w:val", namespaces=NS) == ["3"]
    assert ordered_level.xpath("./w:numFmt/@w:val", namespaces=NS) == [
        "lowerRoman"
    ]
    assert ordered_level.xpath("./w:lvlText/@w:val", namespaces=NS) == ["(%1)"]
    assert ordered_level.xpath("./w:lvlJc/@w:val", namespaces=NS) == ["right"]
    assert ordered_level.xpath("./w:pPr/w:ind/@w:left", namespaces=NS) == ["840"]
    assert ordered_level.xpath("./w:pPr/w:ind/@w:hanging", namespaces=NS) == [
        "240"
    ]
    assert [etree.QName(child).localname for child in ordered_level] == [
        "start",
        "numFmt",
        "lvlText",
        "lvlJc",
        "pPr",
    ]

    fallback_level = ordered_abstract.xpath(
        "./w:lvl[@w:ilvl='8']",
        namespaces=NS,
    )[0]
    assert fallback_level.xpath("./w:numFmt/@w:val", namespaces=NS) == [
        "upperLetter"
    ]
    assert fallback_level.xpath("./w:lvlText/@w:val", namespaces=NS) == ["[%9]"]
    assert fallback_level.xpath("./w:lvlJc/@w:val", namespaces=NS) == ["center"]
    assert fallback_level.xpath("./w:pPr/w:ind/@w:left", namespaces=NS) == [
        "1200"
    ]

    unordered_level = unordered_abstract.xpath(
        "./w:lvl[@w:ilvl='0']",
        namespaces=NS,
    )[0]
    assert unordered_level.xpath("./w:numFmt/@w:val", namespaces=NS) == ["bullet"]
    assert unordered_level.xpath("./w:lvlText/@w:val", namespaces=NS) == ["◆"]
    assert unordered_level.xpath("./w:lvlJc/@w:val", namespaces=NS) == ["center"]
    assert unordered_level.xpath("./w:pPr/w:ind/@w:left", namespaces=NS) == [
        "600"
    ]
    assert unordered_level.xpath(
        "./w:pPr/w:ind/@w:hanging",
        namespaces=NS,
    ) == ["200"]

    styled_paragraph = paragraphs[0]
    assert styled_paragraph.xpath(
        "./w:r/w:rPr/w:rFonts/@w:eastAsia",
        namespaces=NS,
    ) == ["楷体"]
    assert styled_paragraph.xpath(
        "./w:r/w:rPr/w:rFonts/@w:ascii",
        namespaces=NS,
    ) == ["Arial"]
    assert styled_paragraph.xpath("./w:r/w:rPr/w:sz/@w:val", namespaces=NS) == [
        "22"
    ]
    assert styled_paragraph.xpath(
        "./w:r/w:rPr/w:color/@w:val",
        namespaces=NS,
    ) == ["123456"]
    assert styled_paragraph.xpath("./w:r/w:rPr/w:b", namespaces=NS)
    assert styled_paragraph.xpath("./w:r/w:rPr/w:i", namespaces=NS)
    assert styled_paragraph.xpath("./w:pPr/w:jc/@w:val", namespaces=NS) == [
        "both"
    ]
    assert styled_paragraph.xpath(
        "./w:pPr/w:spacing/@w:before",
        namespaces=NS,
    ) == ["120"]
    assert styled_paragraph.xpath(
        "./w:pPr/w:spacing/@w:after",
        namespaces=NS,
    ) == ["160"]
    assert styled_paragraph.xpath(
        "./w:pPr/w:spacing/@w:line",
        namespaces=NS,
    ) == ["360"]
    assert styled_paragraph.xpath(
        "./w:pPr/w:spacing/@w:lineRule",
        namespaces=NS,
    ) == ["exact"]

    abstract_ids = set(
        numbering_xml.xpath("./w:abstractNum/@w:abstractNumId", namespaces=NS)
    )
    referenced_ids = set(
        numbering_xml.xpath("./w:num/w:abstractNumId/@w:val", namespaces=NS)
    )
    assert referenced_ids <= abstract_ids
    children = [etree.QName(child).localname for child in numbering_xml]
    assert max(index for index, name in enumerate(children) if name == "abstractNum") < min(
        index for index, name in enumerate(children) if name == "num"
    )


def test_docx_renderer_creates_real_figures_captions_bookmarks_and_three_line_table(
    tmp_path: Path,
):
    images = tmp_path / "images"
    images.mkdir()
    (images / "model.png").write_bytes(PNG_1X1)
    template = load_template("templates/base/bachelor.yaml")
    template.figure.default_width = LengthSpec.model_validate("40mm")
    template.figure.caption.alignment = "left"
    template.figure.caption.font = FontSpec(east_asia="黑体", latin="Arial")
    template.figure.caption.size = LengthSpec.model_validate("10pt")
    template.table.caption.alignment = "right"
    document = ThesisDocument(
        source_path=tmp_path / "thesis.md",
        blocks=[
            Heading(id="chap:intro", level=1, inlines=[Text(value="绪论")]),
            Figure(
                id="fig:model",
                src="./images/model.png",
                caption_inlines=_text_inlines("系统模型"),
                width="50%",
            ),
            Figure(
                id="fig:default",
                src="./images/model.png",
                caption_inlines=_text_inlines("默认宽度"),
            ),
            _structured_table(
                "tbl:results",
                "实验结果",
                [
                    (True, [("模型", None), ("AUROC", "right")]),
                    (False, [("A", None), ("0.91", "right")]),
                    (False, [("B", None), ("0.94", "right")]),
                ],
            ),
        ],
    )
    output = tmp_path / "figures-tables.docx"

    DocxRenderer().render(compile_document(document, template=template), output)

    document_xml = _xml_part(output, "word/document.xml")
    relationships_xml = _xml_part(output, "word/_rels/document.xml.rels")
    package_parts = list_package_parts(output)
    image_relationships = relationships_xml.xpath(
        "./pr:Relationship[contains(@Type, '/image')]",
        namespaces=REL_NS,
    )
    assert image_relationships
    assert any(part.startswith("word/media/") for part in package_parts)
    assert len(document_xml.xpath(".//w:drawing", namespaces=NS)) == 2
    assert document_xml.xpath(".//wp:extent/@cx", namespaces=NS) == [
        "2790190",
        "1440000",
    ]
    for drawing_paragraph in document_xml.xpath(".//w:p[.//w:drawing]", namespaces=NS):
        assert drawing_paragraph.xpath(
            "./w:pPr/w:spacing/@w:lineRule",
            namespaces=NS,
        ) == ["auto"]
        assert drawing_paragraph.xpath(
            "./w:pPr/w:spacing/@w:line",
            namespaces=NS,
        ) == ["240"]
    assert document_xml.xpath(".//a:blip/@r:embed", namespaces=NS)

    bookmark_names = document_xml.xpath(".//w:bookmarkStart/@w:name", namespaces=NS)
    assert {"tf_fig_model", "tf_fig_default", "tf_tbl_results"} <= set(bookmark_names)
    bookmark_starts = {
        node.get(f"{{{NS['w']}}}id"): node.get(f"{{{NS['w']}}}name")
        for node in document_xml.xpath(".//w:bookmarkStart", namespaces=NS)
    }
    bookmark_ends = set(document_xml.xpath(".//w:bookmarkEnd/@w:id", namespaces=NS))
    assert set(bookmark_starts) <= bookmark_ends

    body = document_xml.find(".//w:body", namespaces=NS)
    assert body is not None
    body_children = list(body)
    drawing_indices = [
        index
        for index, child in enumerate(body_children)
        if child.xpath(".//w:drawing", namespaces=NS)
    ]
    paragraph_text = {
        index: "".join(child.xpath(".//w:t/text()", namespaces=NS))
        for index, child in enumerate(body_children)
        if etree.QName(child).localname == "p"
    }
    assert paragraph_text[drawing_indices[0] + 1] == "图1-1 系统模型"
    assert paragraph_text[drawing_indices[1] + 1] == "图1-2 默认宽度"
    table_index = next(
        index
        for index, child in enumerate(body_children)
        if etree.QName(child).localname == "tbl"
    )
    assert paragraph_text[table_index - 1] == "表1-1 实验结果"

    figure_caption = document_xml.xpath(
        ".//w:p[.//w:bookmarkStart[@w:name='tf_fig_model']]",
        namespaces=NS,
    )[0]
    assert figure_caption.xpath("./w:pPr/w:jc/@w:val", namespaces=NS) == ["left"]
    assert set(
        figure_caption.xpath(".//w:rFonts/@w:eastAsia", namespaces=NS)
    ) == {"黑体"}
    assert set(figure_caption.xpath(".//w:rFonts/@w:ascii", namespaces=NS)) == {
        "Arial"
    }
    assert set(figure_caption.xpath(".//w:sz/@w:val", namespaces=NS)) == {"20"}
    table_caption = document_xml.xpath(
        ".//w:p[.//w:bookmarkStart[@w:name='tf_tbl_results']]",
        namespaces=NS,
    )[0]
    assert table_caption.xpath("./w:pPr/w:jc/@w:val", namespaces=NS) == ["right"]

    table = document_xml.xpath(".//w:tbl", namespaces=NS)[0]
    assert table.xpath(".//w:tr[1]/w:tc//w:t/text()", namespaces=NS) == ["模型", "AUROC"]
    assert table.xpath(".//w:tr[2]/w:tc//w:t/text()", namespaces=NS) == ["A", "0.91"]
    borders = table.xpath("./w:tblPr/w:tblBorders", namespaces=NS)[0]
    assert borders.xpath("./w:top/@w:val", namespaces=NS) == ["single"]
    assert borders.xpath("./w:top/@w:sz", namespaces=NS) == ["12"]
    assert borders.xpath("./w:bottom/@w:val", namespaces=NS) == ["single"]
    assert borders.xpath("./w:bottom/@w:sz", namespaces=NS) == ["12"]
    for edge in ("left", "right", "insideH", "insideV"):
        assert borders.xpath(f"./w:{edge}/@w:val", namespaces=NS) == ["nil"]
    assert table.xpath(
        ".//w:tr[1]/w:tc/w:tcPr/w:tcBorders/w:bottom/@w:val",
        namespaces=NS,
    ) == ["single", "single"]
    assert table.xpath(
        ".//w:tr[1]/w:tc/w:tcPr/w:tcBorders/w:bottom/@w:sz",
        namespaces=NS,
    ) == ["6", "6"]


def test_docx_renderer_honors_configured_three_line_table_widths(tmp_path: Path):
    template = load_template("templates/base/bachelor.yaml")
    template.table.three_line.top_width = LengthSpec.model_validate("2pt")
    template.table.three_line.header_width = LengthSpec.model_validate("0.5pt")
    template.table.three_line.bottom_width = LengthSpec.model_validate("1pt")
    document = ThesisDocument(
        source_path=tmp_path / "thesis.md",
        blocks=[
            _structured_table(
                "tbl:results",
                "实验结果",
                [
                    (True, [("模型", None), ("AUROC", "right")]),
                    (False, [("A", None), ("0.91", "right")]),
                ],
            )
        ],
    )
    output = tmp_path / "configured-three-line.docx"

    DocxRenderer().render(compile_document(document, template=template), output)

    document_xml = _xml_part(output, "word/document.xml")
    table = document_xml.xpath(".//w:tbl", namespaces=NS)[0]
    assert table.xpath(
        "./w:tblPr/w:tblBorders/w:top/@w:sz",
        namespaces=NS,
    ) == ["16"]
    assert table.xpath(
        "./w:tblPr/w:tblBorders/w:bottom/@w:sz",
        namespaces=NS,
    ) == ["8"]
    assert table.xpath(
        ".//w:tr[1]/w:tc/w:tcPr/w:tcBorders/w:bottom/@w:sz",
        namespaces=NS,
    ) == ["4", "4"]


@pytest.mark.parametrize(
    ("width", "expected_size"),
    [("0.25pt", "2"), ("12pt", "96")],
)
def test_docx_renderer_supports_word_border_width_limits(
    tmp_path: Path,
    width: str,
    expected_size: str,
):
    template = load_template("templates/base/bachelor.yaml")
    template.table.three_line.top_width = LengthSpec.model_validate(width)
    document = ThesisDocument(
        source_path=tmp_path / "thesis.md",
        blocks=[
            _structured_table(
                "tbl:limits",
                "线宽边界",
                [
                    (True, [("项目", None), ("值", None)]),
                    (False, [("A", None), ("1", None)]),
                ],
            )
        ],
    )
    output = tmp_path / "three-line-width-limit.docx"

    DocxRenderer().render(compile_document(document, template=template), output)

    document_xml = _xml_part(output, "word/document.xml")
    assert document_xml.xpath(
        ".//w:tbl/w:tblPr/w:tblBorders/w:top/@w:sz",
        namespaces=NS,
    ) == [expected_size]


def test_docx_renderer_honors_non_default_caption_positions(tmp_path: Path):
    images = tmp_path / "images"
    images.mkdir()
    (images / "model.png").write_bytes(PNG_1X1)
    template = load_template("templates/base/bachelor.yaml")
    template.figure.caption.position = "top"
    template.table.caption.position = "bottom"
    document = ThesisDocument(
        source_path=tmp_path / "thesis.md",
        blocks=[
            Figure(
                id="fig:model",
                src="./images/model.png",
                caption_inlines=_text_inlines("系统模型"),
            ),
            _structured_table(
                "tbl:results",
                "实验结果",
                [
                    (True, [("模型", None), ("AUROC", "right")]),
                    (False, [("A", None), ("0.91", "right")]),
                ],
            ),
        ],
    )
    output = tmp_path / "caption-positions.docx"

    DocxRenderer().render(compile_document(document, template=template), output)

    document_xml = _xml_part(output, "word/document.xml")
    body = document_xml.find(".//w:body", namespaces=NS)
    assert body is not None
    children = [
        (
            etree.QName(child).localname,
            "".join(child.xpath(".//w:t/text()", namespaces=NS)),
            bool(child.xpath(".//w:drawing", namespaces=NS)),
        )
        for child in body
        if etree.QName(child).localname != "sectPr"
    ]
    assert children == [
        ("p", "图1-1 系统模型", False),
        ("p", "", True),
        ("tbl", "模型AUROCA0.91", False),
        ("p", "表1-1 实验结果", False),
    ]


def test_docx_renderer_preserves_intrinsic_image_size_without_width_policy(
    tmp_path: Path,
):
    image = tmp_path / "model.png"
    image.write_bytes(PNG_1X1)
    template = load_template("templates/base/bachelor.yaml")
    template.figure.default_width = None
    document = ThesisDocument(
        source_path=tmp_path / "thesis.md",
        blocks=[
            Figure(
                id="fig:model",
                src="./model.png",
                caption_inlines=_text_inlines("系统模型"),
            )
        ],
    )
    output = tmp_path / "intrinsic-width.docx"

    DocxRenderer().render(compile_document(document, template=template), output)

    document_xml = _xml_part(output, "word/document.xml")
    assert document_xml.xpath(".//wp:extent/@cx", namespaces=NS) == ["12700"]


@pytest.mark.parametrize(
    ("style", "expected"),
    [
        ("grid", "single"),
        ("plain", "nil"),
    ],
)
def test_docx_renderer_applies_non_three_line_table_border_policies(
    tmp_path: Path,
    style: str,
    expected: str,
):
    template = load_template("templates/base/bachelor.yaml")
    template.table.style = style
    document = ThesisDocument(
        source_path=tmp_path / "thesis.md",
        blocks=[
            _structured_table(
                "tbl:results",
                "实验结果",
                [
                    (True, [("模型", None), ("AUROC", "right")]),
                    (False, [("A", None), ("0.91", "right")]),
                ],
            )
        ],
    )
    output = tmp_path / f"{style}.docx"

    DocxRenderer().render(compile_document(document, template=template), output)

    document_xml = _xml_part(output, "word/document.xml")
    table = document_xml.xpath(".//w:tbl", namespaces=NS)[0]
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        assert table.xpath(
            f"./w:tblPr/w:tblBorders/w:{edge}/@w:val",
            namespaces=NS,
        ) == [expected]
    assert not table.xpath(".//w:tcBorders", namespaces=NS)


def test_docx_renderer_does_not_create_fake_table_for_empty_rows(tmp_path: Path):
    template = load_template("templates/base/bachelor.yaml")
    document = ThesisDocument(
        source_path=tmp_path / "thesis.md",
        blocks=[_structured_table("tbl:empty", "空表", [])],
    )
    output = tmp_path / "empty-table.docx"

    DocxRenderer().render(compile_document(document, template=template), output)

    document_xml = _xml_part(output, "word/document.xml")
    assert not document_xml.xpath(".//w:tbl", namespaces=NS)
    assert "".join(
        document_xml.xpath(
            ".//w:p[.//w:bookmarkStart[@w:name='tf_tbl_empty']]//w:t/text()",
            namespaces=NS,
        )
    ) == "表1-1 空表"


def test_docx_renderer_creates_real_math_fields_footnotes_and_page_structures(
    tmp_path: Path,
):
    image = tmp_path / "model.png"
    image.write_bytes(PNG_1X1)
    template = load_template("templates/base/bachelor.yaml")
    template.sections = SectionsSpec.model_validate(
        {
            "cover": {
                "start": "new_page",
                "page_number": {"format": "none"},
            },
            "front_matter": {
                "start": "new_page",
                "footer": {"enabled": True},
                "page_number": {"format": "roman-lower"},
            },
            "main": {
                "start": "odd_page",
                "header": {
                    "enabled": True,
                    "text": "基于结构化 Markdown 的论文",
                    "different_first_page": True,
                },
                "footer": {"enabled": True},
                "page_number": {"format": "decimal", "restart": 1},
            },
        }
    )
    document = ThesisDocument(
        source_path=tmp_path / "thesis.md",
        blocks=[
            Heading(
                id="chap:abstract-zh",
                level=1,
                inlines=_text_inlines("摘要"),
            ),
            Paragraph(inlines=[Text(value="摘要正文")]),
            Heading(
                id="chap:introduction",
                level=1,
                inlines=_text_inlines("绪论"),
            ),
            Figure(
                id="fig:model",
                src="./model.png",
                caption_inlines=_text_inlines("系统模型"),
            ),
            _structured_table(
                "tbl:data",
                "实验数据",
                [
                    (True, [("模型", None), ("值", "right")]),
                    (False, [("A", None), ("1", "right")]),
                ],
            ),
            Equation(
                id="eq:loss",
                latex=r"L=-\sum_{i=1}^n y_i \log \hat{y}_i+x_i^2",
                display=True,
            ),
            Paragraph(
                inlines=[
                    Text(value="参见"),
                    CrossReference(target="fig:model"),
                    Text(value="与"),
                    CrossReference(target="eq:loss"),
                    FootnoteReference(label="note"),
                ],
            ),
            FootnoteDefinition(
                label="note",
                inlines=[
                    Text(value="真实脚注正文，参见"),
                    CrossReference(target="fig:model"),
                ],
            ),
        ],
    )
    output = tmp_path / "advanced.docx"

    DocxRenderer().render(compile_document(document, template=template), output)

    package_parts = list_package_parts(output)
    document_xml = _xml_part(output, "word/document.xml")
    settings_xml = _xml_part(output, "word/settings.xml")
    relationships_xml = _xml_part(output, "word/_rels/document.xml.rels")
    content_types_xml = _xml_part(output, "[Content_Types].xml")

    assert document_xml.xpath(".//m:oMath", namespaces=NS)
    assert document_xml.xpath(".//m:nary", namespaces=NS)
    assert document_xml.xpath(".//m:func", namespaces=NS)
    assert document_xml.xpath(".//m:acc", namespaces=NS)
    assert document_xml.xpath(".//m:sSubSup", namespaces=NS)
    assert document_xml.xpath(
        ".//w:p[.//w:bookmarkStart[@w:name='tf_eq_loss']]//m:oMath",
        namespaces=NS,
    )
    equation_paragraph = document_xml.xpath(
        ".//w:p[.//w:bookmarkStart[@w:name='tf_eq_loss']]",
        namespaces=NS,
    )[0]
    equation_children = list(equation_paragraph)
    equation_math = next(
        index
        for index, child in enumerate(equation_children)
        if etree.QName(child).localname == "oMath"
    )
    equation_start = next(
        index
        for index, child in enumerate(equation_children)
        if etree.QName(child).localname == "bookmarkStart"
    )
    equation_end = next(
        index
        for index, child in enumerate(equation_children)
        if etree.QName(child).localname == "bookmarkEnd"
    )
    assert equation_math < equation_start < equation_end
    assert "".join(
        text
        for child in equation_children[equation_start + 1 : equation_end]
        for text in child.xpath(".//w:t/text()", namespaces=NS)
    ) == "(1-1)"
    equation_bookmark_id = equation_children[equation_start].get(f"{{{NS['w']}}}id")
    assert equation_children[equation_end].get(f"{{{NS['w']}}}id") == equation_bookmark_id
    field_codes = [
        "".join(node.itertext()).strip()
        for node in document_xml.xpath(".//w:instrText", namespaces=NS)
    ]
    assert any(code.startswith("SEQ TF_Figure_1") for code in field_codes)
    assert any(code.startswith("SEQ TF_Table_1") for code in field_codes)
    assert any(code.startswith("SEQ TF_Equation_1") for code in field_codes)
    assert "REF tf_fig_model \\h" in field_codes
    assert "REF tf_eq_loss \\h" in field_codes
    assert 'TOC \\o "1-3" \\h \\z \\u' in field_codes
    field_types = document_xml.xpath(".//w:fldChar/@w:fldCharType", namespaces=NS)
    assert field_types.count("begin") == field_types.count("separate")
    assert field_types.count("begin") == field_types.count("end")
    assert settings_xml.xpath("./w:updateFields/@w:val", namespaces=NS) == ["true"]
    for instruction_text in document_xml.xpath(".//w:instrText", namespaces=NS):
        run = instruction_text.getparent()
        container = run.getparent()
        code = "".join(instruction_text.itertext()).strip()
        runs = list(container)
        instruction_index = runs.index(run)
        begin = runs[instruction_index - 1].find("w:fldChar", namespaces=NS)
        separate = runs[instruction_index + 1].find("w:fldChar", namespaces=NS)
        assert begin is not None
        assert begin.get(f"{{{NS['w']}}}fldCharType") == "begin"
        assert begin.get(f"{{{NS['w']}}}dirty") == "true"
        assert separate is not None
        assert separate.get(f"{{{NS['w']}}}fldCharType") == "separate"
        if code.startswith("TOC "):
            # TOC 字段跨段落：begin/separate 在字段段，cached 条目独立成段，
            # end 在最后一个条目段（Word 原生结构）。
            assert any(
                fld.get(f"{{{NS['w']}}}fldCharType") == "end"
                for sibling in container.itersiblings()
                for fld in sibling.iter(f"{{{NS['w']}}}fldChar")
            )
            continue
        assert any(
            candidate.get(f"{{{NS['w']}}}fldCharType") == "end"
            for later_run in runs[instruction_index + 2 :]
            for candidate in later_run.findall("w:fldChar", namespaces=NS)
        )

    figure_caption = document_xml.xpath(
        ".//w:p[.//w:bookmarkStart[@w:name='tf_fig_model']]",
        namespaces=NS,
    )[0]
    children = list(figure_caption)
    bookmark_start = next(
        index
        for index, child in enumerate(children)
        if etree.QName(child).localname == "bookmarkStart"
    )
    bookmark_end = next(
        index
        for index, child in enumerate(children)
        if etree.QName(child).localname == "bookmarkEnd"
    )
    bookmarked_text = "".join(
        text
        for child in children[bookmark_start + 1 : bookmark_end]
        for text in child.xpath(".//w:t/text()", namespaces=NS)
    )
    following_text = "".join(
        text
        for child in children[bookmark_end + 1 :]
        for text in child.xpath(".//w:t/text()", namespaces=NS)
    )
    assert bookmarked_text == "图1-1"
    assert following_text == " 系统模型"

    assert "word/footnotes.xml" in package_parts
    footnotes_xml = _xml_part(output, "word/footnotes.xml")
    assert footnotes_xml.xpath("./w:footnote[@w:id='-1']", namespaces=NS)
    assert footnotes_xml.xpath("./w:footnote[@w:id='0']", namespaces=NS)
    assert footnotes_xml.xpath(
        "./w:footnote[@w:id='1']//w:t[text()='真实脚注正文，参见']",
        namespaces=NS,
    )
    assert footnotes_xml.xpath(
        "./w:footnote[@w:id='1']//w:instrText[text()='REF tf_fig_model \\h']",
        namespaces=NS,
    )
    assert footnotes_xml.xpath(
        "./w:footnote[@w:id='1']//w:fldChar/@w:fldCharType",
        namespaces=NS,
    ) == ["begin", "separate", "end"]
    assert document_xml.xpath(".//w:footnoteReference/@w:id", namespaces=NS) == ["1"]
    assert relationships_xml.xpath(
        "./pr:Relationship[contains(@Type, '/footnotes')]",
        namespaces=REL_NS,
    )
    assert content_types_xml.xpath(
        "./*[local-name()='Override'][@PartName='/word/footnotes.xml']"
    )

    sections = document_xml.xpath(".//w:sectPr", namespaces=NS)
    assert len(sections) == 3
    assert sections[1].xpath("./w:pgNumType/@w:fmt", namespaces=NS) == ["lowerRoman"]
    assert sections[2].xpath("./w:type/@w:val", namespaces=NS) == ["oddPage"]
    assert sections[2].xpath("./w:pgNumType/@w:fmt", namespaces=NS) == ["decimal"]
    assert sections[2].xpath("./w:pgNumType/@w:start", namespaces=NS) == ["1"]
    assert sections[2].xpath("./w:titlePg", namespaces=NS)
    assert sections[1].xpath("./w:footerReference/@r:id", namespaces=NS)
    assert sections[2].xpath("./w:headerReference/@r:id", namespaces=NS)
    assert sections[2].xpath("./w:footerReference/@r:id", namespaces=NS)
    assert relationships_xml.xpath(
        "./pr:Relationship[contains(@Type, '/header')]",
        namespaces=REL_NS,
    )
    assert relationships_xml.xpath(
        "./pr:Relationship[contains(@Type, '/footer')]",
        namespaces=REL_NS,
    )
    header_text = "".join(
        text
        for part in package_parts
        if part.startswith("word/header") and part.endswith(".xml")
        for text in _xml_part(output, part).xpath(".//w:t/text()", namespaces=NS)
    )
    assert "基于结构化 Markdown 的论文" in header_text
    footer_field_codes = [
        "".join(node.itertext()).strip()
        for part in package_parts
        if part.startswith("word/footer") and part.endswith(".xml")
        for node in _xml_part(output, part).xpath(".//w:instrText", namespaces=NS)
    ]
    assert "PAGE" in footer_field_codes
    assert "SECTIONPAGES" in footer_field_codes


def test_docx_renderer_omits_page_fields_when_page_number_format_is_none(
    tmp_path: Path,
):
    template = load_template("templates/base/bachelor.yaml")
    template.sections = SectionsSpec.model_validate(
        {
            "main": {
                "footer": {"enabled": True, "text": "保密文档"},
                "page_number": {"format": "none"},
            }
        }
    )
    document = ThesisDocument(
        source_path=tmp_path / "thesis.md",
        blocks=[
            Heading(
                id="chap:intro",
                level=1,
                inlines=_text_inlines("绪论"),
            )
        ],
    )
    output = tmp_path / "no-page-number.docx"

    DocxRenderer().render(compile_document(document, template=template), output)

    footer_parts = [
        part
        for part in list_package_parts(output)
        if part.startswith("word/footer") and part.endswith(".xml")
    ]
    assert len(footer_parts) == 1
    footer_xml = _xml_part(output, footer_parts[0])
    assert footer_xml.xpath(".//w:t/text()", namespaces=NS) == ["保密文档"]
    assert not footer_xml.xpath(".//w:instrText", namespaces=NS)


def test_docx_renderer_prevents_disabled_section_header_footer_inheritance(
    tmp_path: Path,
):
    template = load_template("templates/base/bachelor.yaml")
    template.sections = SectionsSpec.model_validate(
        {
            "front_matter": {
                "header": {"enabled": True, "text": "FRONT HEADER"},
                "footer": {"enabled": True, "text": "FRONT FOOTER"},
                "page_number": {"format": "roman-lower"},
            },
            "main": {
                "header": {"enabled": False, "text": "DISABLED HEADER TEXT"},
                "footer": {"enabled": False, "text": "DISABLED FOOTER TEXT"},
                "page_number": {"format": "none"},
            },
        }
    )
    document = ThesisDocument(
        source_path=tmp_path / "thesis.md",
        blocks=[
            Heading(
                id="chap:abstract-zh",
                level=1,
                inlines=_text_inlines("摘要"),
            ),
            Heading(
                id="chap:intro",
                level=1,
                inlines=_text_inlines("绪论"),
            ),
        ],
    )
    output = tmp_path / "disabled-header-footer.docx"

    DocxRenderer().render(compile_document(document, template=template), output)

    reloaded = __import__("docx").Document(output)
    assert len(reloaded.sections) == 2
    assert reloaded.sections[1].header.is_linked_to_previous is False
    assert reloaded.sections[1].footer.is_linked_to_previous is False
    assert not any(
        paragraph.text for paragraph in reloaded.sections[1].header.paragraphs
    )
    assert not any(
        paragraph.text for paragraph in reloaded.sections[1].footer.paragraphs
    )

    document_xml = _xml_part(output, "word/document.xml")
    sections = document_xml.xpath(".//w:sectPr", namespaces=NS)
    main_header_id = sections[1].xpath("./w:headerReference/@r:id", namespaces=NS)
    main_footer_id = sections[1].xpath("./w:footerReference/@r:id", namespaces=NS)
    assert main_header_id
    assert main_footer_id
    relationships = _xml_part(output, "word/_rels/document.xml.rels")
    targets = {
        node.get("Id"): node.get("Target")
        for node in relationships.xpath("./pr:Relationship", namespaces=REL_NS)
    }
    main_header = _xml_part(output, f"word/{targets[main_header_id[0]]}")
    main_footer = _xml_part(output, f"word/{targets[main_footer_id[0]]}")
    assert not main_header.xpath(".//w:t[text()='FRONT HEADER']", namespaces=NS)
    assert not main_footer.xpath(".//w:t[text()='FRONT FOOTER']", namespaces=NS)
    assert not main_header.xpath(
        ".//w:t[text()='DISABLED HEADER TEXT']",
        namespaces=NS,
    )
    assert not main_footer.xpath(
        ".//w:t[text()='DISABLED FOOTER TEXT']",
        namespaces=NS,
    )
    assert not main_footer.xpath(".//w:instrText", namespaces=NS)


def test_docx_renderer_writes_page_geometry_and_all_header_footer_variants(
    tmp_path: Path,
):
    template = load_template("templates/base/bachelor.yaml")
    template.page.header_distance = LengthSpec.model_validate("12mm")
    template.page.footer_distance = LengthSpec.model_validate("14mm")
    template.page.document_grid = DocumentGridSpec.model_validate(
        {
            "type": "lines_and_chars",
            "line_pitch": "18pt",
            "char_space": 100,
        }
    )
    template.sections = SectionsSpec.model_validate(
        {
            "main": {
                "header": {
                    "default": {
                        "text": "ODD HEADER",
                        "style": {
                            "font": {"east_asia": "黑体", "latin": "Arial"},
                            "size": "10pt",
                            "alignment": "right",
                            "space_after": "3pt",
                            "snap_to_grid": True,
                        },
                        "bottom_border": {
                            "style": "double",
                            "width": "0.75pt",
                            "color": "336699",
                            "space": "1pt",
                        },
                    },
                    "first": {"text": "FIRST HEADER"},
                    "even": {"text": "EVEN HEADER"},
                },
                "footer": {
                    "default": {
                        "text": "P",
                        "page_number": {
                            "alignment": "right",
                            "page_prefix": "[",
                            "page_suffix": "]",
                            "include_total": False,
                        },
                    },
                    "first": {"enabled": False},
                    "even": {
                        "page_number": {
                            "alignment": "left",
                            "page_prefix": "p",
                            "page_suffix": "",
                            "include_total": True,
                            "separator": "|",
                            "total_prefix": "n",
                            "total_suffix": "",
                        }
                    },
                },
                "page_number": {"format": "decimal", "restart": 3},
            }
        }
    )
    document = ThesisDocument(
        source_path=tmp_path / "thesis.md",
        blocks=[
            Heading(
                id="chap:intro",
                level=1,
                inlines=_text_inlines("绪论"),
            )
        ],
    )
    output = tmp_path / "header-footer-variants.docx"

    DocxRenderer().render(compile_document(document, template=template), output)

    document_xml = _xml_part(output, "word/document.xml")
    settings_xml = _xml_part(output, "word/settings.xml")
    relationships = _xml_part(output, "word/_rels/document.xml.rels")
    section = document_xml.xpath(".//w:sectPr", namespaces=NS)[0]
    margins = section.xpath("./w:pgMar", namespaces=NS)[0]
    assert margins.get(f"{{{NS['w']}}}header") == "680"
    assert margins.get(f"{{{NS['w']}}}footer") == "794"
    assert section.xpath("./w:docGrid/@w:type", namespaces=NS) == [
        "linesAndChars"
    ]
    assert section.xpath("./w:docGrid/@w:linePitch", namespaces=NS) == ["360"]
    assert section.xpath("./w:docGrid/@w:charSpace", namespaces=NS) == ["100"]
    assert section.xpath("./w:pgNumType/@w:fmt", namespaces=NS) == ["decimal"]
    assert section.xpath("./w:pgNumType/@w:start", namespaces=NS) == ["3"]
    assert section.xpath("./w:titlePg", namespaces=NS)
    assert settings_xml.xpath("./w:evenAndOddHeaders", namespaces=NS)
    section_order = [etree.QName(child).localname for child in section]
    assert section_order.index("pgNumType") < section_order.index("cols")
    assert section_order.index("cols") < section_order.index("docGrid")
    settings_order = [etree.QName(child).localname for child in settings_xml]
    assert settings_order.index("evenAndOddHeaders") < settings_order.index(
        "updateFields"
    )

    references = {
        (etree.QName(reference).localname, reference.get(f"{{{NS['w']}}}type")):
        reference.get(f"{{{NS['r']}}}id")
        for reference in section.xpath(
            "./w:headerReference | ./w:footerReference",
            namespaces=NS,
        )
    }
    assert set(references) == {
        ("headerReference", "default"),
        ("headerReference", "first"),
        ("headerReference", "even"),
        ("footerReference", "default"),
        ("footerReference", "first"),
        ("footerReference", "even"),
    }
    targets = {
        relationship.get("Id"): relationship.get("Target")
        for relationship in relationships.xpath(
            "./pr:Relationship",
            namespaces=REL_NS,
        )
    }

    odd_header = _xml_part(
        output,
        f"word/{targets[references[('headerReference', 'default')]]}",
    )
    assert odd_header.xpath(".//w:t/text()", namespaces=NS) == ["ODD HEADER"]
    assert odd_header.xpath(".//w:pPr/w:jc/@w:val", namespaces=NS) == ["right"]
    assert odd_header.xpath(".//w:pPr/w:spacing/@w:after", namespaces=NS) == [
        "60"
    ]
    assert odd_header.xpath(".//w:rPr/w:rFonts/@w:eastAsia", namespaces=NS) == [
        "黑体"
    ]
    assert odd_header.xpath(".//w:rPr/w:rFonts/@w:ascii", namespaces=NS) == [
        "Arial"
    ]
    assert odd_header.xpath(".//w:rPr/w:sz/@w:val", namespaces=NS) == ["20"]
    assert odd_header.xpath(".//w:pBdr/w:bottom/@w:val", namespaces=NS) == [
        "double"
    ]
    assert odd_header.xpath(".//w:pBdr/w:bottom/@w:sz", namespaces=NS) == ["6"]
    assert odd_header.xpath(".//w:pBdr/w:bottom/@w:color", namespaces=NS) == [
        "336699"
    ]
    assert odd_header.xpath(".//w:pBdr/w:bottom/@w:space", namespaces=NS) == [
        "1"
    ]
    paragraph_properties = odd_header.xpath(".//w:pPr", namespaces=NS)[0]
    property_order = [
        etree.QName(child).localname for child in paragraph_properties
    ]
    assert property_order.index("pBdr") < property_order.index("snapToGrid")
    assert property_order.index("snapToGrid") < property_order.index("spacing")
    assert property_order.index("spacing") < property_order.index("jc")

    first_header = _xml_part(
        output,
        f"word/{targets[references[('headerReference', 'first')]]}",
    )
    even_header = _xml_part(
        output,
        f"word/{targets[references[('headerReference', 'even')]]}",
    )
    assert first_header.xpath(".//w:t/text()", namespaces=NS) == ["FIRST HEADER"]
    assert even_header.xpath(".//w:t/text()", namespaces=NS) == ["EVEN HEADER"]

    odd_footer = _xml_part(
        output,
        f"word/{targets[references[('footerReference', 'default')]]}",
    )
    assert odd_footer.xpath(".//w:instrText/text()", namespaces=NS) == ["PAGE"]
    assert odd_footer.xpath(".//w:pPr/w:jc/@w:val", namespaces=NS) == ["right"]
    assert "".join(odd_footer.xpath(".//w:t/text()", namespaces=NS)) == "P [1]"

    first_footer = _xml_part(
        output,
        f"word/{targets[references[('footerReference', 'first')]]}",
    )
    assert not first_footer.xpath(".//w:t | .//w:instrText", namespaces=NS)

    even_footer = _xml_part(
        output,
        f"word/{targets[references[('footerReference', 'even')]]}",
    )
    assert even_footer.xpath(".//w:instrText/text()", namespaces=NS) == [
        "PAGE",
        "SECTIONPAGES",
    ]
    assert even_footer.xpath(".//w:pPr/w:jc/@w:val", namespaces=NS) == ["left"]
    assert "".join(even_footer.xpath(".//w:t/text()", namespaces=NS)) == "p1|n1"


def test_docx_renderer_clears_disabled_variants_in_added_section(
    tmp_path: Path,
):
    template = load_template("templates/base/bachelor.yaml")
    template.page.header_distance = LengthSpec.model_validate("11mm")
    template.page.footer_distance = LengthSpec.model_validate("13mm")
    template.page.document_grid = DocumentGridSpec.model_validate(
        {
            "type": "lines",
            "line_pitch": "20pt",
        }
    )
    enabled_variants = {
        "default": {"text": "DEFAULT"},
        "first": {"text": "FIRST"},
        "even": {"text": "EVEN"},
    }
    disabled_variants = {
        "default": {"enabled": False, "text": "STALE DEFAULT"},
        "first": {"enabled": False, "text": "STALE FIRST"},
        "even": {"enabled": False, "text": "STALE EVEN"},
    }
    template.sections = SectionsSpec.model_validate(
        {
            "front_matter": {
                "header": enabled_variants,
                "footer": enabled_variants,
                "page_number": {"format": "roman-lower"},
            },
            "main": {
                "header": disabled_variants,
                "footer": disabled_variants,
                "page_number": {"format": "none"},
            },
        }
    )
    document = ThesisDocument(
        source_path=tmp_path / "thesis.md",
        blocks=[
            Heading(
                id="chap:abstract-zh",
                level=1,
                inlines=_text_inlines("摘要"),
            ),
            Heading(
                id="chap:intro",
                level=1,
                inlines=_text_inlines("绪论"),
            ),
        ],
    )
    output = tmp_path / "cleared-variants.docx"

    DocxRenderer().render(compile_document(document, template=template), output)

    document_xml = _xml_part(output, "word/document.xml")
    relationships = _xml_part(output, "word/_rels/document.xml.rels")
    sections = document_xml.xpath(".//w:sectPr", namespaces=NS)
    assert len(sections) == 2
    for section in sections:
        margins = section.xpath("./w:pgMar", namespaces=NS)[0]
        assert margins.get(f"{{{NS['w']}}}header") == "624"
        assert margins.get(f"{{{NS['w']}}}footer") == "737"
        assert section.xpath("./w:docGrid/@w:type", namespaces=NS) == ["lines"]
        assert section.xpath("./w:docGrid/@w:linePitch", namespaces=NS) == ["400"]

    main = sections[1]
    references = main.xpath(
        "./w:headerReference | ./w:footerReference",
        namespaces=NS,
    )
    assert {
        (etree.QName(reference).localname, reference.get(f"{{{NS['w']}}}type"))
        for reference in references
    } == {
        ("headerReference", "default"),
        ("headerReference", "first"),
        ("headerReference", "even"),
        ("footerReference", "default"),
        ("footerReference", "first"),
        ("footerReference", "even"),
    }
    targets = {
        relationship.get("Id"): relationship.get("Target")
        for relationship in relationships.xpath(
            "./pr:Relationship",
            namespaces=REL_NS,
        )
    }
    for reference in references:
        relationship_id = reference.get(f"{{{NS['r']}}}id")
        part = _xml_part(
            output,
            f"word/{targets[relationship_id]}",
        )
        assert not part.xpath(".//w:t | .//w:instrText", namespaces=NS)


def test_docx_renderer_uses_current_default_when_even_variant_is_omitted(
    tmp_path: Path,
):
    template = load_template("templates/base/bachelor.yaml")
    template.sections = SectionsSpec.model_validate(
        {
            "front_matter": {
                "header": {
                    "default": {"text": "FRONT DEFAULT"},
                    "even": {"text": "FRONT EVEN"},
                },
                "page_number": {"format": "roman-lower"},
            },
            "main": {
                "header": {
                    "default": {"text": "MAIN DEFAULT"},
                },
                "page_number": {"format": "decimal"},
            },
        }
    )
    document = ThesisDocument(
        source_path=tmp_path / "thesis.md",
        blocks=[
            Heading(
                id="chap:abstract-zh",
                level=1,
                inlines=_text_inlines("摘要"),
            ),
            Heading(
                id="chap:intro",
                level=1,
                inlines=_text_inlines("绪论"),
            ),
        ],
    )
    output = tmp_path / "even-default-fallback.docx"

    DocxRenderer().render(compile_document(document, template=template), output)

    document_xml = _xml_part(output, "word/document.xml")
    relationships = _xml_part(output, "word/_rels/document.xml.rels")
    sections = document_xml.xpath(".//w:sectPr", namespaces=NS)
    main_even_id = sections[1].xpath(
        "./w:headerReference[@w:type='even']/@r:id",
        namespaces=NS,
    )
    assert main_even_id
    targets = {
        relationship.get("Id"): relationship.get("Target")
        for relationship in relationships.xpath(
            "./pr:Relationship",
            namespaces=REL_NS,
        )
    }
    main_even = _xml_part(output, f"word/{targets[main_even_id[0]]}")
    assert main_even.xpath(".//w:t/text()", namespaces=NS) == ["MAIN DEFAULT"]
    assert not main_even.xpath(".//w:t[text()='FRONT EVEN']", namespaces=NS)


def test_docx_renderer_uses_current_default_when_first_variant_is_omitted(
    tmp_path: Path,
):
    template = load_template("templates/base/bachelor.yaml")
    template.sections = SectionsSpec.model_validate(
        {
            "front_matter": {
                "header": {
                    "default": {"text": "FRONT HEADER"},
                    "first": {"text": "FRONT FIRST HEADER"},
                },
                "footer": {
                    "default": {"text": "FRONT FOOTER"},
                    "first": {"text": "FRONT FIRST FOOTER"},
                },
                "page_number": {"format": "roman-lower"},
            },
            "main": {
                "header": {
                    "default": {"text": "MAIN HEADER"},
                    "first": {"text": "MAIN FIRST HEADER"},
                },
                "footer": {
                    "default": {"text": "MAIN FOOTER"},
                },
                "page_number": {"format": "decimal"},
            },
        }
    )
    document = ThesisDocument(
        source_path=tmp_path / "thesis.md",
        blocks=[
            Heading(
                id="chap:abstract-zh",
                level=1,
                inlines=_text_inlines("摘要"),
            ),
            Heading(
                id="chap:intro",
                level=1,
                inlines=_text_inlines("绪论"),
            ),
        ],
    )
    output = tmp_path / "first-default-fallback.docx"

    DocxRenderer().render(compile_document(document, template=template), output)

    document_xml = _xml_part(output, "word/document.xml")
    relationships = _xml_part(output, "word/_rels/document.xml.rels")
    sections = document_xml.xpath(".//w:sectPr", namespaces=NS)
    front_first_footer_id = sections[0].xpath(
        "./w:footerReference[@w:type='first']/@r:id",
        namespaces=NS,
    )
    main_default_footer_id = sections[1].xpath(
        "./w:footerReference[@w:type='default']/@r:id",
        namespaces=NS,
    )
    main_first_footer_id = sections[1].xpath(
        "./w:footerReference[@w:type='first']/@r:id",
        namespaces=NS,
    )
    assert front_first_footer_id
    assert main_default_footer_id
    assert main_first_footer_id
    assert len(
        {
            front_first_footer_id[0],
            main_default_footer_id[0],
            main_first_footer_id[0],
        }
    ) == 3
    targets = {
        relationship.get("Id"): relationship.get("Target")
        for relationship in relationships.xpath(
            "./pr:Relationship",
            namespaces=REL_NS,
        )
    }
    assert len(
        {
            targets[front_first_footer_id[0]],
            targets[main_default_footer_id[0]],
            targets[main_first_footer_id[0]],
        }
    ) == 3
    relationship_types = {
        relationship.get("Id"): relationship.get("Type")
        for relationship in relationships.xpath(
            "./pr:Relationship",
            namespaces=REL_NS,
        )
    }
    assert relationship_types[main_first_footer_id[0]].endswith("/footer")
    front_first_footer = _xml_part(
        output,
        f"word/{targets[front_first_footer_id[0]]}",
    )
    assert front_first_footer.xpath(".//w:t/text()", namespaces=NS) == [
        "FRONT FIRST FOOTER"
    ]
    main_first_footer = _xml_part(
        output,
        f"word/{targets[main_first_footer_id[0]]}",
    )
    assert main_first_footer.xpath(".//w:t/text()", namespaces=NS) == [
        "MAIN FOOTER"
    ]
    assert not main_first_footer.xpath(
        ".//w:t[text()='FRONT FIRST FOOTER']",
        namespaces=NS,
    )


def test_docx_renderer_materializes_initial_first_fallback_with_default_policy(
    tmp_path: Path,
):
    template = load_template("templates/base/bachelor.yaml")
    template.sections = SectionsSpec.model_validate(
        {
            "main": {
                "header": {
                    "default": {
                        "text": "MAIN HEADER",
                        "style": {
                            "font": {"east_asia": "黑体", "latin": "Arial"},
                            "size": "10pt",
                            "space_after": "2pt",
                        },
                        "bottom_border": {
                            "style": "single",
                            "width": "0.5pt",
                        },
                        "page_number": {
                            "alignment": "right",
                            "page_prefix": "(",
                            "page_suffix": ")",
                            "include_total": False,
                        },
                    }
                },
                "footer": {
                    "default": {"text": "MAIN FOOTER"},
                    "first": {"text": "FIRST FOOTER"},
                },
                "page_number": {"format": "decimal"},
            }
        }
    )
    document = ThesisDocument(
        source_path=tmp_path / "thesis.md",
        blocks=[
            Heading(
                id="chap:intro",
                level=1,
                inlines=_text_inlines("绪论"),
            )
        ],
    )
    output = tmp_path / "initial-first-default-policy.docx"

    DocxRenderer().render(compile_document(document, template=template), output)

    reloaded = Document(output)
    assert reloaded.sections[0].first_page_header.is_linked_to_previous is False

    document_xml = _xml_part(output, "word/document.xml")
    relationships = _xml_part(output, "word/_rels/document.xml.rels")
    section = document_xml.xpath(".//w:sectPr", namespaces=NS)[0]
    first_header_id = section.xpath(
        "./w:headerReference[@w:type='first']/@r:id",
        namespaces=NS,
    )
    assert first_header_id
    targets = {
        relationship.get("Id"): relationship.get("Target")
        for relationship in relationships.xpath(
            "./pr:Relationship",
            namespaces=REL_NS,
        )
    }
    first_header = _xml_part(output, f"word/{targets[first_header_id[0]]}")
    assert first_header.xpath(".//w:instrText/text()", namespaces=NS) == ["PAGE"]
    assert "".join(first_header.xpath(".//w:t/text()", namespaces=NS)) == (
        "MAIN HEADER (1)"
    )
    assert first_header.xpath(".//w:pPr/w:jc/@w:val", namespaces=NS) == [
        "right"
    ]
    assert first_header.xpath(".//w:pPr/w:spacing/@w:after", namespaces=NS) == [
        "40"
    ]
    assert first_header.xpath(".//w:pBdr/w:bottom/@w:val", namespaces=NS) == [
        "single"
    ]
    assert first_header.xpath(".//w:pBdr/w:bottom/@w:sz", namespaces=NS) == [
        "4"
    ]
    assert set(
        first_header.xpath(".//w:rPr/w:rFonts/@w:eastAsia", namespaces=NS)
    ) == {"黑体"}
    assert set(
        first_header.xpath(".//w:rPr/w:rFonts/@w:ascii", namespaces=NS)
    ) == {"Arial"}


def test_docx_renderer_materializes_disabled_default_as_blank_first_fallback(
    tmp_path: Path,
):
    template = load_template("templates/base/bachelor.yaml")
    template.sections = SectionsSpec.model_validate(
        {
            "front_matter": {
                "header": {
                    "first": {"text": "FRONT FIRST HEADER"},
                },
                "footer": {
                    "first": {
                        "text": "FRONT FIRST FOOTER",
                        "page_number": {"include_total": True},
                    },
                },
                "page_number": {"format": "roman-lower"},
            },
            "main": {
                "header": {
                    "first": {"text": "MAIN FIRST HEADER"},
                },
                "footer": {
                    "default": {
                        "enabled": False,
                        "text": "DISABLED MAIN FOOTER",
                        "page_number": {"include_total": True},
                    },
                },
                "page_number": {"format": "decimal"},
            },
        }
    )
    document = ThesisDocument(
        source_path=tmp_path / "thesis.md",
        blocks=[
            Heading(
                id="chap:abstract-zh",
                level=1,
                inlines=_text_inlines("摘要"),
            ),
            Heading(
                id="chap:intro",
                level=1,
                inlines=_text_inlines("绪论"),
            ),
        ],
    )
    output = tmp_path / "disabled-first-default-fallback.docx"

    DocxRenderer().render(compile_document(document, template=template), output)

    reloaded = Document(output)
    assert reloaded.sections[1].first_page_footer.is_linked_to_previous is False
    assert not any(
        paragraph.text
        for paragraph in reloaded.sections[1].first_page_footer.paragraphs
    )

    document_xml = _xml_part(output, "word/document.xml")
    relationships = _xml_part(output, "word/_rels/document.xml.rels")
    main = document_xml.xpath(".//w:sectPr", namespaces=NS)[1]
    first_footer_id = main.xpath(
        "./w:footerReference[@w:type='first']/@r:id",
        namespaces=NS,
    )
    assert first_footer_id
    targets = {
        relationship.get("Id"): relationship.get("Target")
        for relationship in relationships.xpath(
            "./pr:Relationship",
            namespaces=REL_NS,
        )
    }
    first_footer = _xml_part(output, f"word/{targets[first_footer_id[0]]}")
    assert not first_footer.xpath(".//w:t | .//w:instrText", namespaces=NS)


def test_clear_header_footer_part_removes_all_blocks_and_relationships(
    tmp_path: Path,
):
    image = tmp_path / "header.png"
    image.write_bytes(PNG_1X1)
    document = Document()
    header = document.sections[0].header
    header.is_linked_to_previous = False
    header.paragraphs[0].add_run("STALE TEXT")
    header.add_table(rows=1, cols=1, width=Inches(1))
    header.add_paragraph().add_run().add_picture(str(image), width=Inches(0.1))
    assert header._element.xpath("./w:tbl")
    assert any(
        relationship.reltype.endswith("/image")
        for relationship in header.part.rels.values()
    )

    sections_module._clear_part(header)

    assert [etree.QName(child).localname for child in header._element] == ["p"]
    assert not header._element.xpath(".//w:t | .//w:tbl | .//w:drawing")
    assert not any(
        relationship.reltype.endswith("/image")
        for relationship in header.part.rels.values()
    )


def test_reference_field_runs_centralize_ref_instruction():
    reference = ReferenceRun(
        target_id="fig:model",
        bookmark="tf_fig_model",
        display_text="图1-1",
    )

    runs = fields_module.reference_field_runs(reference)

    instructions = [
        "".join(node.itertext())
        for run in runs
        for node in run.xpath("./w:instrText")
    ]
    assert instructions == ["REF tf_fig_model \\h"]


def test_docx_renderer_writes_resolved_body_footnote_and_bibliography_text(
    tmp_path: Path,
):
    fixture = Path(__file__).parent / "fixtures" / "bibliography" / "gbt7714-v1.bib"
    database = LocalBibTeXLoader().load(fixture)
    body_citation = Citation(
        keys=["doe2024", "smith2025"],
        raw="[@doe2024; @smith2025]",
    )
    footnote_citation = Citation(keys=["smith2025"], raw="[@smith2025]")
    document = ThesisDocument(
        source_path=tmp_path / "thesis.md",
        blocks=[
            Paragraph(
                inlines=[
                    Text(value="正文引用"),
                    body_citation,
                    FootnoteReference(label="note"),
                ],
            ),
            FootnoteDefinition(
                label="note",
                inlines=[Text(value="脚注引用"), footnote_citation],
            ),
            BibliographyBlock(),
        ],
        citations=[body_citation, footnote_citation],
    )
    output = tmp_path / "bibliography.docx"

    plan = compile_document(
        document,
        template=load_template("templates/base/bachelor.yaml"),
        bibliography_database=database,
        citation_formatter=Gbt7714Formatter(),
    )
    DocxRenderer().render(plan, output)

    document_xml = _xml_part(output, "word/document.xml")
    footnotes_xml = _xml_part(output, "word/footnotes.xml")
    body_text = "".join(document_xml.xpath(".//w:body//w:t/text()", namespaces=NS))
    footnote_text = "".join(
        footnotes_xml.xpath("./w:footnote[@w:id='1']//w:t/text()", namespaces=NS)
    )
    paragraphs = [
        "".join(paragraph.xpath(".//w:t/text()", namespaces=NS))
        for paragraph in document_xml.xpath(".//w:body/w:p", namespaces=NS)
    ]

    assert "正文引用[1,2]" in body_text
    assert "[@doe2024" not in body_text
    assert "脚注引用[2]" in footnote_text
    assert "[@smith2025]" not in footnote_text
    assert paragraphs[-2:] == [
        "[1] DOE J. Structured Academic Documents[M]. Beijing: Example Press, 2024.",
        (
            "[2] SMITH J, WANG L. Deterministic Thesis Compilation[J]. "
            "Journal of Document Engineering, 2025, 12(3): 101-118. "
            "DOI:10.1000/tf.2025.001."
        ),
    ]


@pytest.mark.parametrize(
    ("presentation", "expected_superscript"),
    [
        (None, False),
        ("inline", False),
        ("superscript", True),
    ],
)
def test_docx_renderer_applies_citation_presentation_only_to_citation_runs(
    tmp_path: Path,
    presentation: str | None,
    expected_superscript: bool,
):
    template = load_template("templates/base/bachelor.yaml")
    assert template.citation is not None
    if presentation is None:
        template.citation = None
    else:
        template.citation.presentation = presentation
    output = tmp_path / f"citation-{presentation or 'omitted'}.docx"

    DocxRenderer().render(
        RenderPlan(
            nodes=[
                ParagraphInstruction(
                    text="引用[1,2, p. 12]。",
                    inlines=(
                        TextRun("引用"),
                        CitationRun(
                            keys=("doe2024", "smith2025"),
                            ordinals=(1, 2),
                            locator="p. 12",
                            raw="[@doe2024; @smith2025, p. 12]",
                            text="[1,2, p. 12]",
                        ),
                        TextRun("。"),
                    ),
                )
            ],
            template=template,
        ),
        output,
    )

    document_xml = _xml_part(output, "word/document.xml")
    citation_run = document_xml.xpath(
        ".//w:r[w:t[text()='[1,2, p. 12]']]",
        namespaces=NS,
    )[0]
    text_runs = document_xml.xpath(
        ".//w:r[w:t[text()='引用' or text()='。']]",
        namespaces=NS,
    )

    assert "".join(document_xml.xpath(".//w:body//w:t/text()", namespaces=NS)) == (
        "引用[1,2, p. 12]。"
    )
    assert bool(
        citation_run.xpath(
            "./w:rPr/w:vertAlign[@w:val='superscript']",
            namespaces=NS,
        )
    ) is expected_superscript
    assert all(
        not run.xpath("./w:rPr/w:vertAlign", namespaces=NS)
        for run in text_runs
    )


def test_docx_renderer_applies_superscript_to_body_and_footnote_citations(
    tmp_path: Path,
):
    fixture = Path(__file__).parent / "fixtures" / "bibliography" / "gbt7714-v1.bib"
    database = LocalBibTeXLoader().load(fixture)
    body_citation = Citation(keys=["doe2024"], raw="[@doe2024]")
    footnote_citation = Citation(keys=["smith2025"], raw="[@smith2025]")
    document = ThesisDocument(
        source_path=tmp_path / "thesis.md",
        blocks=[
            Paragraph(
                inlines=[
                    Text(value="正文引用"),
                    body_citation,
                    FootnoteReference(label="note"),
                ],
            ),
            FootnoteDefinition(
                label="note",
                inlines=[Text(value="脚注引用"), footnote_citation],
            ),
        ],
        citations=[body_citation, footnote_citation],
    )
    template = load_template("templates/base/bachelor.yaml")
    assert template.citation is not None
    template.citation.presentation = "superscript"
    output = tmp_path / "citation-footnote-superscript.docx"

    DocxRenderer().render(
        compile_document(
            document,
            template=template,
            bibliography_database=database,
            citation_formatter=Gbt7714Formatter(),
        ),
        output,
    )

    document_xml = _xml_part(output, "word/document.xml")
    footnotes_xml = _xml_part(output, "word/footnotes.xml")
    assert document_xml.xpath(
        ".//w:r[w:t[text()='[1]']]/w:rPr/w:vertAlign/@w:val",
        namespaces=NS,
    ) == ["superscript"]
    assert footnotes_xml.xpath(
        ".//w:footnote[@w:id='1']//w:r[w:t[text()='[2]']]"
        "/w:rPr/w:vertAlign/@w:val",
        namespaces=NS,
    ) == ["superscript"]


def test_docx_renderer_applies_bibliography_title_and_entry_policy_xml(
    tmp_path: Path,
):
    fixture = Path(__file__).parent / "fixtures" / "bibliography" / "gbt7714-v1.bib"
    database = LocalBibTeXLoader().load(fixture)
    citation = Citation(
        keys=["doe2024", "smith2025"],
        locator="p. 12",
        raw="[@doe2024; @smith2025, p. 12]",
    )
    document = ThesisDocument(
        source_path=tmp_path / "thesis.md",
        blocks=[
            Paragraph(inlines=[Text(value="引用"), citation]),
            Heading(
                id="references",
                level=1,
                inlines=_text_inlines("参考文献"),
            ),
            BibliographyBlock(),
        ],
        citations=[citation],
    )
    template = load_template("templates/base/bachelor.yaml")
    template.bibliography = BibliographySpec(
        title=ParagraphStyleSpec(
            font=FontSpec(east_asia="黑体", latin="Arial"),
            size="16pt",
            bold=True,
            alignment="center",
            space_before="12pt",
            space_after="6pt",
        ),
        entry=ParagraphStyleSpec(
            font=FontSpec(east_asia="宋体", latin="Times New Roman"),
            size="10.5pt",
            left_indent="2em",
            hanging_indent="2em",
            space_before="6pt",
            space_after="0pt",
            line_spacing={"type": "fixed", "value": "20pt"},
        ),
    )
    output = tmp_path / "bibliography-policy.docx"

    DocxRenderer().render(
        compile_document(
            document,
            template=template,
            bibliography_database=database,
            citation_formatter=Gbt7714Formatter(),
        ),
        output,
    )

    styles_xml = _xml_part(output, "word/styles.xml")
    document_xml = _xml_part(output, "word/document.xml")
    title_style = styles_xml.xpath(
        ".//w:style[@w:styleId='TFBibliographyTitle']",
        namespaces=NS,
    )[0]
    entry_style = styles_xml.xpath(
        ".//w:style[@w:styleId='TFBibliographyEntry']",
        namespaces=NS,
    )[0]

    assert title_style.xpath("./w:basedOn/@w:val", namespaces=NS) == ["Heading1"]
    assert title_style.xpath("./w:rPr/w:rFonts/@w:eastAsia", namespaces=NS) == [
        "黑体"
    ]
    assert title_style.xpath("./w:rPr/w:rFonts/@w:ascii", namespaces=NS) == ["Arial"]
    assert title_style.xpath("./w:rPr/w:sz/@w:val", namespaces=NS) == ["32"]
    assert title_style.xpath("./w:rPr/w:b", namespaces=NS)
    assert title_style.xpath("./w:pPr/w:jc/@w:val", namespaces=NS) == ["center"]
    assert title_style.xpath("./w:pPr/w:spacing/@w:before", namespaces=NS) == ["240"]
    assert title_style.xpath("./w:pPr/w:spacing/@w:after", namespaces=NS) == ["120"]
    assert entry_style.xpath("./w:basedOn/@w:val", namespaces=NS) == ["Normal"]
    assert entry_style.xpath("./w:rPr/w:rFonts/@w:eastAsia", namespaces=NS) == [
        "宋体"
    ]
    assert entry_style.xpath("./w:rPr/w:rFonts/@w:ascii", namespaces=NS) == [
        "Times New Roman"
    ]
    assert entry_style.xpath("./w:rPr/w:sz/@w:val", namespaces=NS) == ["21"]
    assert entry_style.xpath("./w:pPr/w:ind/@w:left", namespaces=NS) == ["420"]
    assert entry_style.xpath("./w:pPr/w:ind/@w:hanging", namespaces=NS) == ["420"]
    assert entry_style.xpath("./w:pPr/w:spacing/@w:before", namespaces=NS) == ["120"]
    assert entry_style.xpath("./w:pPr/w:spacing/@w:after", namespaces=NS) == ["0"]
    assert entry_style.xpath("./w:pPr/w:spacing/@w:line", namespaces=NS) == ["400"]
    assert entry_style.xpath(
        "./w:pPr/w:spacing/@w:lineRule",
        namespaces=NS,
    ) == ["exact"]

    assert document_xml.xpath(
        ".//w:p[.//w:t[text()='参考文献']]/w:pPr/w:pStyle/@w:val",
        namespaces=NS,
    ) == ["TFBibliographyTitle"]
    bibliography_paragraphs = document_xml.xpath(
        ".//w:p[w:pPr/w:pStyle[@w:val='TFBibliographyEntry']]",
        namespaces=NS,
    )
    assert [
        "".join(paragraph.xpath(".//w:t/text()", namespaces=NS))
        for paragraph in bibliography_paragraphs
    ] == [
        "[1] DOE J. Structured Academic Documents[M]. Beijing: Example Press, 2024.",
        (
            "[2] SMITH J, WANG L. Deterministic Thesis Compilation[J]. "
            "Journal of Document Engineering, 2025, 12(3): 101-118. "
            "DOI:10.1000/tf.2025.001."
        ),
    ]


@pytest.mark.parametrize("error_type", [AttributeError, ValueError])
def test_docx_renderer_wraps_private_api_failures_with_capability_context(
    tmp_path: Path,
    monkeypatch,
    error_type,
):
    document = ThesisDocument(
        source_path=tmp_path / "thesis.md",
        blocks=[Equation(id="eq:broken", latex="x", display=True)],
    )
    plan = compile_document(
        document,
        template=load_template("templates/base/bachelor.yaml"),
    )

    def fail_render(*args, **kwargs):
        raise error_type("missing private member")

    monkeypatch.setattr(renderer_module, "render_equation", fail_render)

    with pytest.raises(DocxRenderError, match="equation.*missing private member"):
        DocxRenderer().render(plan, tmp_path / "broken.docx")


# ---------- ADR-0003：扩展 LaTeX 子集的 OMML 结构断言 ----------


def _render_equation_document_xml(tmp_path: Path, equation_id: str, latex: str):
    template = load_template("templates/base/bachelor.yaml")
    document = ThesisDocument(
        source_path=tmp_path / "thesis.md",
        blocks=[
            Heading(
                id="chap:math",
                level=1,
                inlines=_text_inlines("公式"),
            ),
            Equation(id=equation_id, latex=latex, display=True),
        ],
    )
    output = tmp_path / f"{equation_id.replace(':', '_')}.docx"
    DocxRenderer().render(compile_document(document, template=template), output)
    return _xml_part(output, "word/document.xml")


def _equation_omath(document_xml, equation_id: str):
    bookmark = f"tf_{equation_id.replace(':', '_')}"
    matches = document_xml.xpath(
        f".//w:p[.//w:bookmarkStart[@w:name='{bookmark}']]//m:oMath",
        namespaces=NS,
    )
    assert len(matches) == 1
    return matches[0]


def test_docx_renderer_pmatrix_omml_structure(tmp_path: Path):
    document_xml = _render_equation_document_xml(
        tmp_path, "eq:pmx", r"A = \begin{pmatrix} a & b \\ c & d \end{pmatrix}"
    )
    omath = _equation_omath(document_xml, "eq:pmx")

    matrices = omath.xpath(
        "./m:d[m:dPr/m:begChr[@m:val='(']]"
        "[m:dPr/m:endChr[@m:val=')']]"
        "[m:dPr/m:grow]/m:e/m:m",
        namespaces=NS,
    )
    assert len(matrices) == 1
    matrix = matrices[0]
    rows = matrix.xpath("./m:mr", namespaces=NS)
    assert len(rows) == 2
    for row, expected in zip(rows, (("a", "b"), ("c", "d")), strict=True):
        cells = row.xpath("./m:e", namespaces=NS)
        assert [
            "".join(cell.xpath(".//m:t/text()", namespaces=NS)) for cell in cells
        ] == list(expected)
    columns = matrix.xpath(
        "./m:mPr/m:mcs/m:mc/m:mcPr/m:mcJc[@m:val='center']",
        namespaces=NS,
    )
    assert len(columns) == 2


def test_docx_renderer_vmatrix_uses_vertical_bar_delimiters(tmp_path: Path):
    document_xml = _render_equation_document_xml(
        tmp_path,
        "eq:vmx",
        r"\det A = \begin{vmatrix} a & b \\ c & d \end{vmatrix} = ad - bc",
    )
    omath = _equation_omath(document_xml, "eq:vmx")

    assert omath.xpath(
        "./m:d[m:dPr/m:begChr[@m:val='|']][m:dPr/m:endChr[@m:val='|']]/m:e/m:m",
        namespaces=NS,
    )
    assert omath.xpath(".//m:func[m:fName//m:t='det']", namespaces=NS)


def test_docx_renderer_cases_omml_structure(tmp_path: Path):
    document_xml = _render_equation_document_xml(
        tmp_path,
        "eq:cases",
        r"f(x) = \begin{cases} x^2, & x \geq 0 \\ -x, & x < 0 \end{cases}",
    )
    omath = _equation_omath(document_xml, "eq:cases")

    matrices = omath.xpath(
        "./m:d[m:dPr/m:begChr[@m:val='{']][m:dPr/m:endChr[@m:val='']]/m:e/m:m",
        namespaces=NS,
    )
    assert len(matrices) == 1
    matrix = matrices[0]
    rows = matrix.xpath("./m:mr", namespaces=NS)
    assert len(rows) == 2
    assert all(len(row.xpath("./m:e", namespaces=NS)) == 2 for row in rows)
    assert matrix.xpath(
        "./m:mPr/m:mcs/m:mc/m:mcPr/m:mcJc[@m:val='left']",
        namespaces=NS,
    )
    first_row_cells = rows[0].xpath("./m:e", namespaces=NS)
    assert first_row_cells[0].xpath("./m:sSup[m:sup//m:t='2']", namespaces=NS)
    assert (
        "".join(first_row_cells[1].xpath(".//m:t/text()", namespaces=NS)) == "x≥0"
    )


def test_docx_renderer_aligned_omml_structure(tmp_path: Path):
    document_xml = _render_equation_document_xml(
        tmp_path, "eq:aligned", r"\begin{aligned} y &= ax + b \\ &= c \end{aligned}"
    )
    omath = _equation_omath(document_xml, "eq:aligned")

    arrays = omath.xpath("./m:eqArr", namespaces=NS)
    assert len(arrays) == 1
    rows = arrays[0].xpath("./m:e", namespaces=NS)
    assert len(rows) == 2
    # 对齐点以 OMML 行内 & 记号表达（与 pandoc/texmath 产物同一约定）
    assert rows[0].xpath("./m:r[m:t='&']", namespaces=NS)
    assert rows[1].xpath("./m:r[m:t='&']", namespaces=NS)
    assert "".join(rows[0].xpath(".//m:t/text()", namespaces=NS)) == "y&=ax+b"
    assert "".join(rows[1].xpath(".//m:t/text()", namespaces=NS)) == "&=c"


def test_docx_renderer_left_right_norm_omml_structure(tmp_path: Path):
    document_xml = _render_equation_document_xml(
        tmp_path, "eq:norm", r"\left\| x \right\|_2 = \sqrt{\sum_{i=1}^{n} x_i^2}"
    )
    omath = _equation_omath(document_xml, "eq:norm")

    delimiters = omath.xpath(
        "./m:sSub[m:sub//m:t='2']/m:e/m:d"
        "[m:dPr/m:begChr[@m:val='‖']][m:dPr/m:endChr[@m:val='‖']][m:dPr/m:grow]",
        namespaces=NS,
    )
    assert len(delimiters) == 1


def test_docx_renderer_left_right_fraction_omml_structure(tmp_path: Path):
    document_xml = _render_equation_document_xml(
        tmp_path, "eq:grow", r"\left( \frac{a+b}{c} \right)^2"
    )
    omath = _equation_omath(document_xml, "eq:grow")

    fractions = omath.xpath(
        "./m:sSup[m:sup//m:t='2']/m:e/m:d"
        "[m:dPr/m:begChr[@m:val='(']][m:dPr/m:endChr[@m:val=')']]/m:e/m:f",
        namespaces=NS,
    )
    assert len(fractions) == 1


def test_docx_renderer_lim_uses_limlow(tmp_path: Path):
    document_xml = _render_equation_document_xml(
        tmp_path, "eq:lime", r"\lim_{n \to \infty} (1 + \frac{1}{n})^n = e"
    )
    omath = _equation_omath(document_xml, "eq:lime")

    limits = omath.xpath("./m:limLow", namespaces=NS)
    assert len(limits) == 1
    limit = limits[0]
    assert limit.xpath(
        "./m:e/m:r[m:rPr/m:sty[@m:val='p']][m:t='lim']",
        namespaces=NS,
    )
    lower = limit.xpath("./m:lim", namespaces=NS)
    assert len(lower) == 1
    assert "".join(lower[0].xpath(".//m:t/text()", namespaces=NS)) == "n→∞"


def test_docx_renderer_lim_upper_uses_limupp(tmp_path: Path):
    document_xml = _render_equation_document_xml(
        tmp_path, "eq:limstar", r"\lim^{*} f(x)"
    )
    omath = _equation_omath(document_xml, "eq:limstar")

    uppers = omath.xpath("./m:limUpp", namespaces=NS)
    assert len(uppers) == 1
    upper = uppers[0]
    assert upper.xpath(
        "./m:e/m:r[m:rPr/m:sty[@m:val='p']][m:t='lim']",
        namespaces=NS,
    )
    limit = upper.xpath("./m:lim", namespaces=NS)
    assert len(limit) == 1
    assert "".join(limit[0].xpath(".//m:t/text()", namespaces=NS)) == "*"


def test_docx_renderer_int_and_prod_use_nary(tmp_path: Path):
    document_xml = _render_equation_document_xml(
        tmp_path, "eq:int", r"\int_{a}^{b} f(x) dx = F(b) - F(a)"
    )
    omath = _equation_omath(document_xml, "eq:int")

    integrals = omath.xpath("./m:nary[m:naryPr/m:chr[@m:val='∫']]", namespaces=NS)
    assert len(integrals) == 1
    integral = integrals[0]
    assert integral.xpath("./m:sub[.//m:t='a']", namespaces=NS)
    assert integral.xpath("./m:sup[.//m:t='b']", namespaces=NS)
    assert integral.xpath("./m:e", namespaces=NS)

    document_xml = _render_equation_document_xml(
        tmp_path, "eq:prod", r"L(\theta) = \prod_{i=1}^{n} f(x_i; \theta)"
    )
    omath = _equation_omath(document_xml, "eq:prod")
    products = omath.xpath("./m:nary[m:naryPr/m:chr[@m:val='∏']]", namespaces=NS)
    assert len(products) == 1
    lower = products[0].xpath("./m:sub", namespaces=NS)
    assert len(lower) == 1
    assert "".join(lower[0].xpath(".//m:t/text()", namespaces=NS)) == "i=1"
    assert products[0].xpath("./m:sup[.//m:t='n']", namespaces=NS)


def test_docx_renderer_binom_uses_nobar_fraction_in_delimiters(tmp_path: Path):
    document_xml = _render_equation_document_xml(
        tmp_path, "eq:binom", r"\binom{n}{k} = \frac{n!}{k!(n-k)!}"
    )
    omath = _equation_omath(document_xml, "eq:binom")

    fractions = omath.xpath(
        "./m:d[m:dPr/m:begChr[@m:val='(']][m:dPr/m:endChr[@m:val=')']]"
        "/m:e/m:f[m:fPr/m:type[@m:val='noBar']]",
        namespaces=NS,
    )
    assert len(fractions) == 1
    assert fractions[0].xpath("./m:num[.//m:t='n']", namespaces=NS)
    assert fractions[0].xpath("./m:den[.//m:t='k']", namespaces=NS)


def test_docx_renderer_mathrm_and_text_omml_structure(tmp_path: Path):
    document_xml = _render_equation_document_xml(
        tmp_path,
        "eq:text",
        r"F = ma, \text{其中 } m \text{ 为质量}, \mathrm{MSE}",
    )
    omath = _equation_omath(document_xml, "eq:text")

    assert omath.xpath(".//m:r[m:rPr/m:nor][m:t='其中 ']", namespaces=NS)
    assert omath.xpath(".//m:r[m:rPr/m:nor][m:t=' 为质量']", namespaces=NS)
    assert omath.xpath(
        ".//m:r[not(m:rPr/m:nor)][m:rPr/m:sty[@m:val='p']][m:t='MSE']",
        namespaces=NS,
    )


def test_docx_renderer_extended_accents_omml(tmp_path: Path):
    document_xml = _render_equation_document_xml(
        tmp_path, "eq:accent", r"\vec F = m \ddot a + \tilde x"
    )
    omath = _equation_omath(document_xml, "eq:accent")

    characters = omath.xpath(".//m:acc/m:accPr/m:chr/@m:val", namespaces=NS)
    assert characters == ["→", "¨", "~"]


def test_docx_renderer_function_name_scripts_omml(tmp_path: Path):
    document_xml = _render_equation_document_xml(
        tmp_path, "eq:funcsub", r"\sin^2 \theta + \log_2 n = 1"
    )
    omath = _equation_omath(document_xml, "eq:funcsub")

    assert omath.xpath(
        "./m:sSup[m:sup//m:t='2']/m:e/m:func[m:fName//m:t='sin']",
        namespaces=NS,
    )
    assert omath.xpath(
        "./m:sSub[m:sub//m:t='2']/m:e/m:func[m:fName//m:t='log']",
        namespaces=NS,
    )


def test_docx_renderer_function_argument_keeps_parentheses_inside(tmp_path: Path):
    document_xml = _render_equation_document_xml(
        tmp_path, "eq:entropy", r"H(X) = -\sum_{i=1}^{n} p(x_i) \log p(x_i)"
    )
    omath = _equation_omath(document_xml, "eq:entropy")

    functions = omath.xpath(".//m:func[m:fName//m:t='log']", namespaces=NS)
    assert len(functions) == 1
    texts = functions[0].xpath("./m:e//m:t/text()", namespaces=NS)
    assert "".join(texts) == "p(xi)"


def test_docx_renderer_bare_line_break_raises_explicitly(tmp_path: Path):
    with pytest.raises(MathSyntaxError, match="Bare"):
        _render_equation_document_xml(
            tmp_path, "eq:broken", r"y = a + b \\ z = c + d"
        )
