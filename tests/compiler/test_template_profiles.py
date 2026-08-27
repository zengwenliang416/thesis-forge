from pathlib import Path

import pytest
from docx import Document

from docforge.application import preview_service
from docforge.core.compiler import MissingRequiredBindingError, compile_document
from docforge.core.model import ForgeDocument, Heading, Text
from docforge.core.render_plan import CoverInstruction
from docforge.renderers.docx import DocxRenderer
from docforge.templates import load_template

ROOT = Path(__file__).resolve().parents[2]
FIXTURES = ROOT / "tests" / "fixtures"


def _cover(project_name: str) -> CoverInstruction:
    preview = preview_service(FIXTURES / project_name / "document.md")
    assert preview.plan is not None, preview.issues
    return next(
        node for node in preview.plan.nodes if isinstance(node, CoverInstruction)
    )


def _visible_text(path: Path) -> str:
    document = Document(path)
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def test_general_and_academic_projects_compile_to_typed_binding_plans() -> None:
    general = dict(_cover("docforge-general").bindings)
    academic = dict(_cover("docforge-academic").bindings)

    assert general["metadata.title.zh"] == "DocForge 通用文档"
    assert general["metadata.organization"] == "示例科技有限公司"
    assert general["metadata.date"] == "2026-08-26"
    assert not any(path.startswith("academic.") for path in general)
    assert academic["metadata.title.zh"] == "DocForge 学术文档"
    assert academic["academic.student.name"] == "张三"
    assert academic["academic.degree.name"] == "工学硕士"


def test_docforge_standard_docx_contains_no_fabricated_academic_content(
    tmp_path: Path,
) -> None:
    preview = preview_service(FIXTURES / "docforge-general" / "document.md")
    assert preview.plan is not None, preview.issues
    output = tmp_path / "document.docx"

    DocxRenderer().render(preview.plan, output)
    text = _visible_text(output)

    assert "DocForge 通用文档" in text
    assert "示例科技有限公司" in text
    assert "张三" in text
    for forbidden in (
        "硕士学位论文",
        "本科毕业论文",
        "研究生：",
        "学号：",
        "指导教师：",
        "导师职称：",
        "学位类别：",
        "学科专业：",
    ):
        assert forbidden not in text


def test_absent_optional_common_metadata_omits_labels_and_fallbacks(
    tmp_path: Path,
) -> None:
    template = load_template(ROOT / "templates" / "base" / "docforge-standard.yaml")
    document = ForgeDocument(
        source_path=Path("/tmp/document.md"),
        metadata={
            "metadata": {
                "title": {"zh": "只有标题的文档"},
                "authors": [],
                "keywords": [],
            }
        },
        blocks=[
            Heading(
                id="chap:introduction",
                level=1,
                inlines=[Text(value="正文")],
            )
        ],
    )
    plan = compile_document(document, template=template)
    cover = next(node for node in plan.nodes if isinstance(node, CoverInstruction))
    assert cover.bindings == (("metadata.title.zh", "只有标题的文档"),)
    output = tmp_path / "minimal.docx"

    DocxRenderer().render(plan, output)
    text = _visible_text(output)

    assert "只有标题的文档" in text
    assert "作者：" not in text
    assert "组织：" not in text
    assert "版本：" not in text
    assert "日期：" not in text
    assert "关键词：" not in text


def test_compile_document_rejects_missing_required_academic_bindings() -> None:
    template = load_template(
        ROOT / "templates" / "schools" / "example-university" / "2026.yaml"
    )
    document = ForgeDocument(
        source_path=Path("/tmp/document.md"),
        metadata={
            "metadata": {
                "title": {"zh": "缺少学术资料"},
            }
        },
        blocks=[Heading(level=1, inlines=[Text(value="正文")])],
    )

    with pytest.raises(
        MissingRequiredBindingError,
        match="Required template bindings are missing",
    ) as error:
        compile_document(document, template=template)

    assert error.value.paths == (
        "academic.institution.name",
        "academic.institution.department",
        "academic.degree.name",
        "academic.degree.major",
        "academic.student.name",
        "academic.student.id",
        "academic.advisor.name",
        "academic.completion.date",
    )
