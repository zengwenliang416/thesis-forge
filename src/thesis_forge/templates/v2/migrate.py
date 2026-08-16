"""v0.3 单 YAML → Template Package v2 迁移（SCHEMA §8，决策 D-7）。

输入 v0.3 模板（`templates/model.py` 的 `load_template`），输出目录形态 v2
包骨架：`template.yaml`（可自动映射字段，`schema_version: 2`）+
`reference.docx`（样式由 v0.3 字段编程注入，工程化自 SPIKE 路线①
`spikes/phase0/docx-template/build_reference.py`）+ `provenance.yaml` 骨架 +
`README.md` / `CHANGELOG.md` / `fixtures/minimal/` 骨架。

逐字段三态台账（§8.1 映射表逐行覆盖）写入 `migration-report.json`：
`migrated`（含落点路径）、`manual-required`（含原因与建议操作）、
`dropped`（含理由，仅允许无语义损失——本实现仅用于「显式值与 v2 默认值
一致」的字段）。

行为约束（§8.2）：目标目录非空时拒绝覆盖（`force=True` 显式开启）；迁移
幂等，可重复执行；产物立即跑 L1–L3，error 反映在返回的 `lint_report`。
"""

from __future__ import annotations

import json
import re
from dataclasses import asdict, dataclass, field
from datetime import UTC, date, datetime
from pathlib import Path
from typing import Any

import yaml

from thesis_forge import __version__ as _HOST_VERSION
from thesis_forge.templates.model import (
    CaptionSpec,
    LengthSpec,
    NumberingSpec,
    ThesisTemplate,
    load_template,
)

from . import lint as _lint
from .schema import TEMPLATE_ID_RE

MIGRATED = "migrated"
MANUAL_REQUIRED = "manual-required"
DROPPED = "dropped"
LEDGER_STATUSES = (MIGRATED, MANUAL_REQUIRED, DROPPED)

GENERATED_VERSION = "0.1.0"


class MigrateError(ValueError):
    """迁移前置条件失败（输入不可读 / 输出目录非空等）。"""


@dataclass(frozen=True, slots=True)
class LedgerEntry:
    field: str
    status: str  # migrated / manual-required / dropped
    target: str | None = None
    reason: str | None = None
    suggestion: str | None = None


@dataclass(frozen=True, slots=True)
class MigrateReport:
    source: Path
    output: Path
    entries: tuple[LedgerEntry, ...]
    lint_report: _lint.LintReport
    reference_styles: tuple[str, ...] = field(default_factory=tuple)

    @property
    def summary(self) -> dict[str, int]:
        return {
            status: sum(entry.status == status for entry in self.entries)
            for status in LEDGER_STATUSES
        }

    def to_dict(self) -> dict[str, Any]:
        return {
            "source": str(self.source),
            "output": str(self.output),
            "summary": self.summary,
            "entries": [asdict(entry) for entry in self.entries],
            "lint": {
                "levels_run": list(self.lint_report.levels_run),
                "errors": self.lint_report.errors,
                "warnings": self.lint_report.warnings,
                "issues": [asdict(issue) for issue in self.lint_report.issues],
            },
        }


def _length(value: LengthSpec | None) -> str | None:
    return None if value is None else str(value)


# ---------------------------------------------------------------------------
# template.yaml 字段映射（§8.1）
# ---------------------------------------------------------------------------


def _sanitize_id(raw: str) -> tuple[str, bool]:
    """v0.3 id → §3.1 正则；不符合时改写（小写、非法字符折叠为 -）。"""
    if TEMPLATE_ID_RE.fullmatch(raw):
        return raw, False
    rewritten = re.sub(r"[^a-z0-9.-]+", "-", raw.lower())
    rewritten = re.sub(r"\.{2,}", ".", rewritten).strip("-.")
    rewritten = re.sub(r"-{2,}", "-", rewritten)
    if not rewritten or not TEMPLATE_ID_RE.fullmatch(rewritten):
        rewritten = "migrated-template"
    return rewritten, True


def _map_numbering(
    numbering: NumberingSpec,
    *,
    kind: str,
    ledger: list[LedgerEntry],
) -> dict[str, Any]:
    """v0.3 numbering（mode/separator）→ §3.12 scope/enabled/separator。"""
    data: dict[str, Any] = {}
    if numbering.mode == "none":
        data["enabled"] = False
        ledger.append(
            LedgerEntry(
                field=f"{kind}.numbering.mode",
                status=MIGRATED,
                target=f"template.yaml#numbering.{kind}.enabled",
                reason="mode: none → enabled: false（§3.12 C-1b）",
            )
        )
    else:
        data["scope"] = numbering.mode
        ledger.append(
            LedgerEntry(
                field=f"{kind}.numbering.mode",
                status=MIGRATED,
                target=f"template.yaml#numbering.{kind}.scope",
            )
        )
    if numbering.separator == "-":
        ledger.append(
            LedgerEntry(
                field=f"{kind}.numbering.separator",
                status=DROPPED,
                reason="与 v2 默认值 '-' 一致，省略（无语义损失）",
            )
        )
    else:
        data["separator"] = numbering.separator
        ledger.append(
            LedgerEntry(
                field=f"{kind}.numbering.separator",
                status=MIGRATED,
                target=f"template.yaml#numbering.{kind}.separator",
            )
        )
    return data


def _map_caption_prefix(
    caption: CaptionSpec,
    *,
    kind: str,
    default_prefix: str,
    ledger: list[LedgerEntry],
) -> str | None:
    if caption.prefix == default_prefix:
        ledger.append(
            LedgerEntry(
                field=f"{kind}.caption.prefix",
                status=DROPPED,
                reason=f"与 v2 默认值 {default_prefix!r} 一致，省略（无语义损失）",
            )
        )
        return None
    ledger.append(
        LedgerEntry(
            field=f"{kind}.caption.prefix",
            status=MIGRATED,
            target=f"template.yaml#numbering.{kind}.caption_prefix",
        )
    )
    return caption.prefix


def _build_template_data(
    template: ThesisTemplate,
    raw: dict[str, Any],
    ledger: list[LedgerEntry],
) -> dict[str, Any]:
    data: dict[str, Any] = {}

    # header（§3.1）
    package_id, rewritten = _sanitize_id(template.id)
    ledger.append(
        LedgerEntry(
            field="id",
            status=MIGRATED,
            target="template.yaml#id",
            reason="id 不符合 §3.1 正则，已改写" if rewritten else None,
        )
    )
    ledger.append(LedgerEntry(field="name", status=MIGRATED, target="template.yaml#name"))
    data.update(
        {
            "schema_version": 2,
            "id": package_id,
            "version": GENERATED_VERSION,
            "name": template.name,
            "language": "zh-CN",
            "status": "draft",
        }
    )
    ledger.append(
        LedgerEntry(
            field="year",
            status=MANUAL_REQUIRED,
            reason="v2 header 无 year 落点（§8.1）",
            suggestion="录入 provenance.yaml school.official_document.version"
            "（骨架已按 year 预填，请核对）",
        )
    )

    host_major = int(_HOST_VERSION.split(".")[0])
    host_minor = _HOST_VERSION.split(".")[1] if "." in _HOST_VERSION else "0"
    data["compatibility"] = {
        "thesisforge": f">={host_major}.{host_minor},<{host_major + 1}.0",
        "document_types": ["bachelor_thesis"],
        "target_apps": {"word": "primary"},
    }
    ledger.append(
        LedgerEntry(
            field="compatibility.document_types",
            status=MANUAL_REQUIRED,
            reason="v0.3 无文档类型字段，骨架默认 bachelor_thesis",
            suggestion="按实际学位类型调整 compatibility.document_types",
        )
    )

    # page（§3.5）：margin.left/right → inner/outer + mirror_margins: false
    page = template.page
    page_data: dict[str, Any] = {
        "size": page.size,
        "orientation": page.orientation,
        "margin": {
            "top": str(page.margin.top),
            "bottom": str(page.margin.bottom),
            "inner": str(page.margin.left),
            "outer": str(page.margin.right),
        },
        "mirror_margins": False,
    }
    if page.header_distance is not None:
        page_data["header_distance"] = str(page.header_distance)
    if page.footer_distance is not None:
        page_data["footer_distance"] = str(page.footer_distance)
    if page.document_grid is not None:
        grid: dict[str, Any] = {"type": page.document_grid.type}
        if page.document_grid.line_pitch is not None:
            grid["line_pitch"] = str(page.document_grid.line_pitch)
        if page.document_grid.char_space is not None:
            grid["char_space"] = page.document_grid.char_space
        page_data["document_grid"] = grid
    data["page"] = page_data
    ledger.append(
        LedgerEntry(
            field="page.size/orientation",
            status=MIGRATED,
            target="template.yaml#page",
        )
    )
    ledger.append(
        LedgerEntry(
            field="page.margin.left/right",
            status=MIGRATED,
            target="template.yaml#page.margin.inner/outer",
            reason="v0.3 无镜像边距：inner≡left、outer≡right + mirror_margins: false（§3.5）",
        )
    )
    if page.header_distance is not None or page.document_grid is not None:
        ledger.append(
            LedgerEntry(
                field="page.header_distance/footer_distance/document_grid",
                status=MIGRATED,
                target="template.yaml#page",
            )
        )

    # fonts（§3.6）
    fonts: dict[str, Any] = {
        "body": {
            "east_asia": template.body.font.east_asia,
            "latin": template.body.font.latin,
        }
    }
    heading_font = template.heading.level1.font
    if heading_font is not None:
        fonts["heading"] = {
            "east_asia": heading_font.east_asia,
            "latin": heading_font.latin,
        }
    data["fonts"] = fonts
    ledger.append(
        LedgerEntry(
            field="body.font",
            status=MIGRATED,
            target="template.yaml#fonts.body + reference.docx#TF Body",
        )
    )

    # styles（§3.7）token 映射：reference.docx 由本工具生成同名 TF 样式。
    # headings level2–4 的 style 缺省回落为同级 token，故 1–4 级必须全部声明。
    paragraph_tokens: dict[str, str] = {
        "body": "TF Body",
        "body_first": "TF Body First",
        "abstract": "TF Abstract",
    }
    character_tokens: dict[str, str] = {"code": "TF Code Char"}
    if template.bibliography is not None or template.citation is not None:
        paragraph_tokens["bibliography"] = "TF Bibliography"
    if template.figure is not None:
        paragraph_tokens["caption_figure"] = "TF Figure Caption"
    if template.table is not None:
        paragraph_tokens["caption_table"] = "TF Table Caption"
    if template.equation is not None:
        paragraph_tokens["equation"] = "TF Equation"
        character_tokens["equation_inline"] = "TF Equation Inline"
    data["styles"] = {
        "paragraph": paragraph_tokens,
        "heading": {
            "1": "TF Heading 1",
            "2": "TF Heading 2",
            "3": "TF Heading 3",
            "4": "TF Heading 4",
        },
        "character": character_tokens,
    }

    # body（§3.8）：白名单字段进 YAML，其余由 reference.docx TF Body 承载
    body = template.body
    body_data: dict[str, Any] = {"alignment": body.alignment}
    if body.first_line_indent is not None:
        body_data["first_line_indent"] = str(body.first_line_indent)
    if body.line_spacing is not None:
        spacing: dict[str, Any] = {"type": body.line_spacing.type}
        value = body.line_spacing.value
        if isinstance(value, LengthSpec):
            spacing["value"] = str(value)
        elif value is not None:
            spacing["value"] = value
        body_data["line_spacing"] = spacing
    spacing_data = {
        key: _length(getattr(body, f"space_{key}"))
        for key in ("before", "after")
        if getattr(body, f"space_{key}") is not None
    }
    if spacing_data:
        body_data["spacing"] = spacing_data
    if body.widow_control is not None:
        body_data["widow_control"] = body.widow_control
    data["body"] = body_data
    ledger.append(
        LedgerEntry(
            field="body.alignment/first_line_indent/line_spacing/spacing/widow_control",
            status=MIGRATED,
            target="template.yaml#body",
            reason="§4.1 白名单字段（B 类）",
        )
    )
    ledger.append(
        LedgerEntry(
            field="body.size/color/bold 等样式属性",
            status=MIGRATED,
            target="reference.docx#TF Body",
            reason="§4.1 C 类：纯样式由 token 样式承载",
        )
    )

    # headings（§3.9）：白名单（page_break_before/keep_with_next）进 YAML
    headings_data: dict[str, Any] = {}
    for level in (1, 2, 3):
        spec = getattr(template.heading, f"level{level}")
        if spec is None:
            continue
        entry: dict[str, Any] = {}
        if spec.page_break_before:
            entry["page_break_before"] = True
        if spec.keep_with_next is not None:
            entry["keep_with_next"] = spec.keep_with_next
        if entry:
            headings_data[str(level)] = entry
        ledger.append(
            LedgerEntry(
                field=f"heading.level{level}",
                status=MIGRATED,
                target=f"reference.docx#TF Heading {level}"
                + (" + template.yaml#headings" if entry else ""),
                reason="样式属性 → reference.docx；白名单字段 → §3.9",
            )
        )
    if headings_data:
        data["headings"] = headings_data
    ledger.append(
        LedgerEntry(
            field="heading.level4",
            status=MANUAL_REQUIRED,
            reason="v0.3 无级别 4；TF Heading 4 样式复用级别 3，v2 级别 4 默认关闭编号（§3.9）",
            suggestion="人工核对 reference.docx 的 TF Heading 4 样式",
        )
    )

    # regions（§3.10）：顺序骨架 + 可直迁标题
    order = ["cover", "abstract_zh", "abstract_en", "toc", "main"]
    if template.bibliography is not None or template.citation is not None:
        order.append("bibliography")
    order.append("acknowledgements")
    regions_data: dict[str, Any] = {"order": order}
    data["regions"] = regions_data

    # cover / list / semantic_styles / toc（无 YAML 落点或部件化）
    ledger.append(
        LedgerEntry(
            field="cover.items",
            status=MANUAL_REQUIRED,
            reason="cover items 无 YAML 落点（§8.1：layouts/cover.yaml blocks 或 shell.docx）"
            + ("；v0.3 使用默认 11 项封面字段" if "cover" not in raw else ""),
            suggestion="人工编写 layouts/cover.yaml（§3.20）或 shell.docx 封面",
        )
    )
    ledger.append(
        LedgerEntry(
            field="list.ordered/unordered",
            status=MANUAL_REQUIRED,
            reason="偏差记录 C-11：v2 无 lists 节，由 reference.docx numbering base 承载"
            + ("（v0.3 为默认列表几何）" if "list" not in raw else ""),
            suggestion="人工核对列表编号/缩进渲染效果",
        )
    )
    if template.semantic_styles is not None and "semantic_styles" in raw:
        ledger.append(
            LedgerEntry(
                field="semantic_styles.*",
                status=MIGRATED,
                target="reference.docx#TF Abstract 等",
                reason="v2 YAML 不承载纯样式（§4.1 C 类）；region 标题请人工核对",
            )
        )
    if template.toc is not None:
        toc_data: dict[str, Any] = {}
        for level in (1, 2, 3):
            level_spec = template.toc.for_level(level)
            if level_spec is None:
                continue
            level_data: dict[str, Any] = {}
            if level_spec.leader != "dots":
                level_data["leader"] = level_spec.leader
            if level_spec.page_number_tab is not None:
                level_data["page_number_tab"] = str(level_spec.page_number_tab)
            if level_data:
                toc_data.setdefault("levels", {})[str(level)] = level_data
        if toc_data:
            data["toc"] = toc_data
        ledger.append(
            LedgerEntry(
                field="toc.level1-3.leader/page_number_tab",
                status=MIGRATED,
                target="template.yaml#toc.levels",
            )
        )
        ledger.append(
            LedgerEntry(
                field="toc.title 与其余 toc 样式",
                status=MANUAL_REQUIRED,
                reason="v0.3 toc.title 为样式对象无文本；其余样式 → reference.docx TOC N",
                suggestion="在 regions.toc.title 填写目录标题；核对 TOC N 样式",
            )
        )
    if template.bibliography is not None:
        ledger.append(
            LedgerEntry(
                field="bibliography.title/entry",
                status=MANUAL_REQUIRED,
                reason="entry 样式 → reference.docx TF Bibliography；title 无文本落点",
                suggestion="在 regions.bibliography.title 填写参考文献标题",
            )
        )

    # figure / table / equation（§3.12–§3.15）
    if template.figure is not None:
        figure = template.figure
        data["numbering"] = data.get("numbering", {})
        data["numbering"]["figure"] = _map_numbering(
            figure.numbering, kind="figure", ledger=ledger
        )
        prefix = _map_caption_prefix(
            figure.caption, kind="figure", default_prefix="图", ledger=ledger
        )
        if prefix is not None:
            data["numbering"]["figure"]["caption_prefix"] = prefix
        figures_data: dict[str, Any] = {"caption": {"position": figure.caption.position}}
        if figure.default_width is not None:
            figures_data["default_width"] = str(figure.default_width)
        data["figures"] = figures_data
        ledger.append(
            LedgerEntry(
                field="figure.caption.position/default_width",
                status=MIGRATED,
                target="template.yaml#figures",
            )
        )
        ledger.append(
            LedgerEntry(
                field="figure.caption.font/size/alignment",
                status=MIGRATED,
                target="reference.docx#TF Figure Caption",
            )
        )
    if template.table is not None:
        table = template.table
        data.setdefault("numbering", {})["table"] = _map_numbering(
            table.numbering, kind="table", ledger=ledger
        )
        # NumberingSpec.table 的 TF_TABLE 默认值挂在字段 default_factory 上；
        # 显式声明 numbering.table 时必须自带 sequence_name，否则回落为
        # CaptionNumberingSpec 类默认 TF_FIGURE 而撞上 figure。
        data["numbering"]["table"]["sequence_name"] = "TF_TABLE"
        prefix = _map_caption_prefix(
            table.caption, kind="table", default_prefix="表", ledger=ledger
        )
        if prefix is not None:
            data["numbering"]["table"]["caption_prefix"] = prefix
        borders: dict[str, str]
        if table.style == "three_line":
            borders = {
                "top": str(table.three_line.top_width),
                "header_bottom": str(table.three_line.header_width),
                "bottom": str(table.three_line.bottom_width),
                "inside_vertical": "none",
                "inside_horizontal": "none",
            }
            ledger.append(
                LedgerEntry(
                    field="table.three_line.*",
                    status=MIGRATED,
                    target="template.yaml#tables.styles.three_line.borders",
                )
            )
        elif table.style == "grid":
            borders = {
                "top": "0.5pt",
                "header_bottom": "0.5pt",
                "bottom": "0.5pt",
                "inside_vertical": "0.5pt",
                "inside_horizontal": "0.5pt",
            }
            ledger.append(
                LedgerEntry(
                    field="table.style: grid",
                    status=MANUAL_REQUIRED,
                    reason="grid 展开为自定义样式键（§3.14），线宽按 0.5pt 预填",
                    suggestion="人工核对 tables.styles.grid 线宽",
                )
            )
        else:  # plain
            borders = {
                "top": "none",
                "header_bottom": "none",
                "bottom": "none",
                "inside_vertical": "none",
                "inside_horizontal": "none",
            }
            ledger.append(
                LedgerEntry(
                    field="table.style: plain",
                    status=MANUAL_REQUIRED,
                    reason="plain 展开为自定义样式键（§3.14），边框全部置 none",
                    suggestion="人工核对 tables.styles.plain",
                )
            )
        data["tables"] = {
            "default_style": table.style,
            "caption": {"position": table.caption.position},
            "styles": {table.style: {"borders": borders}},
        }
        ledger.append(
            LedgerEntry(
                field="table.caption.position",
                status=MIGRATED,
                target="template.yaml#tables.caption.position",
            )
        )
        ledger.append(
            LedgerEntry(
                field="table.caption.font/size/alignment",
                status=MIGRATED,
                target="reference.docx#TF Table Caption",
            )
        )
    if template.equation is not None:
        data.setdefault("numbering", {})["equation"] = _map_numbering(
            template.equation.numbering, kind="equation", ledger=ledger
        )
        data["equations"] = {"alignment": template.equation.alignment}
        ledger.append(
            LedgerEntry(
                field="equation.alignment",
                status=MIGRATED,
                target="template.yaml#equations.alignment",
            )
        )
    if not data.get("numbering"):
        data.pop("numbering", None)

    # sections（§3.11）
    sections_data: dict[str, Any] = {}
    for key in ("cover", "front_matter", "main"):
        section = getattr(template.sections, key)
        if section is None:
            continue
        section_data: dict[str, Any] = {"start": section.start}
        page_number = section.page_number
        if page_number.format == "none":
            section_data["page_number"] = {"display": False}
        else:
            number_data: dict[str, Any] = {"display": True, "format": page_number.format}
            if page_number.restart is not None:
                number_data["restart"] = page_number.restart
            section_data["page_number"] = number_data
        sections_data[key] = section_data
        ledger.append(
            LedgerEntry(
                field=f"sections.{key}.start/page_number",
                status=MIGRATED,
                target=f"template.yaml#sections.{key}",
                reason="format: none → display: false（§3.11）"
                if page_number.format == "none"
                else None,
            )
        )
        for part in ("header", "footer"):
            part_spec = getattr(section, part)
            declared = part_spec.enabled or any(
                getattr(part_spec, variant) is not None
                and getattr(part_spec, variant).enabled
                for variant in ("default", "first", "even")
            )
            if declared:
                ledger.append(
                    LedgerEntry(
                        field=f"sections.{key}.{part}",
                        status=MANUAL_REQUIRED,
                        reason="偏差记录 C-12：页眉页脚内容部件化迁移；默认部件文字/页码域"
                        "已注入 reference.docx（PAGE 域 + cached result）",
                        suggestion="人工核对 reference.docx 页眉页脚部件与 "
                        f"sections.{key}.header_footer 引用",
                    )
                )
    if sections_data:
        data["sections"] = sections_data

    # citation → bibliography（§3.19）
    if template.citation is not None:
        bibliography_data: dict[str, Any] = {
            "style_file": "citations/style.csl",
            "presentation": template.citation.presentation,
        }
        data["bibliography"] = bibliography_data
        ledger.append(
            LedgerEntry(
                field="citation.style",
                status=MANUAL_REQUIRED,
                reason=f"工具未内置 CSL 注册表，无法解析 {template.citation.style!r}",
                suggestion="骨架已生成最小占位 citations/style.csl，请替换为对应 "
                "CSL 1.0.1 文件，并在 provenance.yaml 记录其哈希与许可",
            )
        )
        ledger.append(
            LedgerEntry(
                field="citation.presentation",
                status=MIGRATED,
                target="template.yaml#bibliography.presentation",
            )
        )

    return data


# ---------------------------------------------------------------------------
# reference.docx 生成（工程化自 spike 路线① build_reference.py）
# ---------------------------------------------------------------------------


def _build_reference_docx(template: ThesisTemplate, path: Path) -> list[str]:
    """以 v0.3 字段编程注入样式/页面设置/页眉页脚，返回生成的 TF 样式名清单。

    实现思路来源：spikes/phase0/docx-template/build_reference.py（路线①实证），
    此处工程化为按 v0.3 模板参数驱动（不绑定具体学校）。
    """
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from lxml import etree

    from thesis_forge.renderers.docx.document import configure_section_geometry
    from thesis_forge.renderers.docx.styles import apply_paragraph_style
    from thesis_forge.renderers.docx.units import to_docx_length
    from thesis_forge.templates.model import ParagraphStyleSpec

    def add_paragraph_style(document, name: str, base=None):
        style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = base if base is not None else document.styles["Normal"]
        style.quick_style = True
        return style

    def set_default_fonts(document, *, east_asia: str, latin: str, size_pt: float) -> None:
        styles_element = document.styles.element
        doc_defaults = styles_element.find(qn("w:docDefaults"))
        if doc_defaults is None:
            doc_defaults = OxmlElement("w:docDefaults")
            styles_element.insert(0, doc_defaults)
        r_pr_default = doc_defaults.find(qn("w:rPrDefault"))
        if r_pr_default is None:
            r_pr_default = OxmlElement("w:rPrDefault")
            doc_defaults.append(r_pr_default)
        r_pr = r_pr_default.find(qn("w:rPr"))
        if r_pr is None:
            r_pr = OxmlElement("w:rPr")
            r_pr_default.append(r_pr)
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.insert(0, r_fonts)
        for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
            r_fonts.attrib.pop(qn(f"w:{attr}"), None)
        r_fonts.set(qn("w:ascii"), latin)
        r_fonts.set(qn("w:hAnsi"), latin)
        r_fonts.set(qn("w:eastAsia"), east_asia)
        r_fonts.set(qn("w:cs"), latin)
        sz = r_pr.find(qn("w:sz"))
        if sz is None:
            sz = OxmlElement("w:sz")
            r_pr.append(sz)
        sz.set(qn("w:val"), str(int(size_pt * 2)))

    def declare_font_table(document, font_names: list[str]) -> None:
        part = next(
            part
            for part in document.part.package.parts
            if str(part.partname) == "/word/fontTable.xml"
        )
        root = etree.fromstring(part.blob)
        existing = {el.get(qn("w:name")) for el in root.findall(qn("w:font"))}
        for name in font_names:
            if name in existing:
                continue
            font = OxmlElement("w:font")
            font.set(qn("w:name"), name)
            family = OxmlElement("w:family")
            family.set(qn("w:val"), "roman")
            font.append(family)
            charset = OxmlElement("w:charset")
            charset.set(qn("w:val"), "86")  # GB2312，中文字体常规声明
            font.append(charset)
            root.append(font)
        part._blob = etree.tostring(
            root, xml_declaration=True, encoding="UTF-8", standalone=True
        )

    def build_header_footer(document, *, header_text: str | None, page_field: bool) -> None:
        section = document.sections[0]
        if header_text:
            header_paragraph = section.header.paragraphs[0]
            header_paragraph.style = document.styles["Header"]
            header_paragraph.add_run(header_text)
            p_pr = header_paragraph._p.get_or_add_pPr()
            p_bdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")  # 0.5pt，单位 1/8pt
            bottom.set(qn("w:color"), "auto")
            bottom.set(qn("w:space"), "1")
            p_bdr.append(bottom)
            p_pr.append(p_bdr)
        if page_field:
            footer_paragraph = section.footer.paragraphs[0]
            footer_paragraph.style = document.styles["Footer"]
            begin = OxmlElement("w:r")
            fld_begin = OxmlElement("w:fldChar")
            fld_begin.set(qn("w:fldCharType"), "begin")
            begin.append(fld_begin)
            separate = OxmlElement("w:r")
            fld_separate = OxmlElement("w:fldChar")
            fld_separate.set(qn("w:fldCharType"), "separate")
            separate.append(fld_separate)
            cached = OxmlElement("w:r")
            text = OxmlElement("w:t")
            text.text = "1"  # cached result（AGENTS.md §1.5：真域 + 缓存值）
            cached.append(text)
            instr = OxmlElement("w:r")
            instr_text = OxmlElement("w:instrText")
            instr_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            instr_text.text = " PAGE "
            instr.append(instr_text)
            end = OxmlElement("w:r")
            fld_end = OxmlElement("w:fldChar")
            fld_end.set(qn("w:fldCharType"), "end")
            end.append(fld_end)
            footer_paragraph._p.extend((begin, instr, separate, cached, end))

    document = Document()  # 以 python-docx 默认包为骨架，注入模板样式
    configure_section_geometry(document.sections[0], template.page)

    body_size_pt = float(to_docx_length(template.body.size).pt)
    set_default_fonts(
        document,
        east_asia=template.body.font.east_asia,
        latin=template.body.font.latin,
        size_pt=body_size_pt,
    )

    normal = document.styles["Normal"]
    apply_paragraph_style(
        normal,
        ParagraphStyleSpec(font=template.body.font, size=template.body.size),
    )

    body_style = add_paragraph_style(document, "TF Body")
    apply_paragraph_style(body_style, template.body)
    add_paragraph_style(document, "TF Body First", base=body_style)

    add_paragraph_style(document, "TF Abstract")
    abstract_source = template.semantic_styles.abstract_zh
    abstract_body = (
        abstract_source.body
        if abstract_source is not None and abstract_source.body is not None
        else template.body
    )
    apply_paragraph_style(document.styles["TF Abstract"], abstract_body)

    heading_fallback = template.heading.level3 or template.heading.level1
    for level in (1, 2, 3, 4):
        spec = getattr(template.heading, f"level{level}", None) or heading_fallback
        style = add_paragraph_style(document, f"TF Heading {level}")
        # L3 outline-level-mismatch：heading token 样式必须带 outline level；
        # v0.3 未声明时按标题级别补齐（0 起）。
        if spec.outline_level is None:
            spec = spec.model_copy(update={"outline_level": level - 1})
        apply_paragraph_style(
            style,
            spec,
            fallback_font=template.body.font,
            fallback_size=template.body.size,
        )

    if template.bibliography is not None or template.citation is not None:
        style = add_paragraph_style(document, "TF Bibliography")
        entry_spec = (
            template.bibliography.entry if template.bibliography is not None else None
        )
        apply_paragraph_style(
            style,
            entry_spec or ParagraphStyleSpec(),
            fallback_font=template.body.font,
            fallback_size=template.body.size,
        )

    if template.figure is not None:
        caption = template.figure.caption
        apply_paragraph_style(
            add_paragraph_style(document, "TF Figure Caption"),
            ParagraphStyleSpec(
                font=caption.font, size=caption.size, alignment=caption.alignment
            ),
            fallback_font=template.body.font,
            fallback_size=template.body.size,
        )
    if template.table is not None:
        caption = template.table.caption
        apply_paragraph_style(
            add_paragraph_style(document, "TF Table Caption"),
            ParagraphStyleSpec(
                font=caption.font, size=caption.size, alignment=caption.alignment
            ),
            fallback_font=template.body.font,
            fallback_size=template.body.size,
        )
    if template.equation is not None:
        apply_paragraph_style(
            add_paragraph_style(document, "TF Equation"),
            ParagraphStyleSpec(
                font=template.body.font,
                size=template.body.size,
                alignment=template.equation.alignment,
                line_spacing=template.body.line_spacing,
            ),
        )
        inline = document.styles.add_style("TF Equation Inline", WD_STYLE_TYPE.CHARACTER)
        inline.font.name = template.body.font.latin

    code_char = document.styles.add_style("TF Code Char", WD_STYLE_TYPE.CHARACTER)
    code_char.font.name = "Consolas"

    font_names = [template.body.font.east_asia, template.body.font.latin]
    if template.heading.level1.font is not None:
        font_names.append(template.heading.level1.font.east_asia)
    declare_font_table(document, list(dict.fromkeys(font_names)))

    # 页眉页脚部件（偏差记录 C-12）：取首个启用 default 变体的文字/页码域
    header_text: str | None = None
    page_field = False
    for key in ("cover", "front_matter", "main"):
        section = getattr(template.sections, key)
        if section is None:
            continue
        default_header = section.header.default
        if (
            header_text is None
            and default_header is not None
            and default_header.enabled
            and default_header.text
        ):
            header_text = default_header.text
        default_footer = section.footer.default
        footer_page = (
            default_footer is not None
            and default_footer.enabled
            and default_footer.page_number is not None
        )
        if footer_page or section.page_number.format != "none":
            page_field = True
    if header_text or page_field:
        build_header_footer(document, header_text=header_text, page_field=page_field)

    document.save(path)
    return sorted(
        style.name
        for style in document.styles
        if style.name and style.name.startswith("TF ")
    )


# ---------------------------------------------------------------------------
# 骨架文件
# ---------------------------------------------------------------------------


def _provenance_yaml(template: ThesisTemplate, *, today: date) -> str:
    year = str(template.year).strip() or "TODO"
    return f"""# 迁移生成的溯源骨架（SCHEMA §3.21）；TODO 项为人工必填（R-024）。
school:
  name: TODO（请填写学校/机构全称）
  official_document:
    title: TODO（请填写学校官方论文规范文件名）
    version: {json.dumps(year, ensure_ascii=False)}  # 由 v0.3 year 预填，请核对（§8.1）
    source_type: manual
maintainers:
  - name: TODO
    contact: TODO
licenses:
  template_code: TODO
  school_assets: TODO
review:
  last_verified: {today.isoformat()}
  verified_with:
    - TODO（例如：Word 365 / WPS / LibreOffice 版本 + 验证方式）
notes: >
  本包由 `thesisforge template migrate` 从 v0.3 模板 {template.id} 迁移生成；
  逐字段台账见 migration-report.json。
"""


_README_TEMPLATE = """# {name}

由 `thesisforge template migrate` 从 v0.3 模板 `{legacy_id}` 迁移生成的
Template Package v2 骨架（schema_version 2）。

## 使用说明

```bash
thesisforge template lint <本目录>
thesisforge template pack <本目录> -o dist/{package_id}-{version}.tftpl
```

## 已知限制

以下字段无法自动迁移，需人工核对（逐字段台账见 `migration-report.json`）：

{manual_items}
"""


_FIXTURE_MD = """---
thesis:
  title: 最小验证文档
author:
  name: 示例作者
---

# 第一章 绪论 {#chap:intro}

正文段落，用于 fixtures/minimal 冒烟。
"""


# L1/L3 要求 bibliography.style_file 声明的 CSL 存在；真实样式文件为
# manual-required（台账），此处生成最小合法 CSL 1.0.1 占位。
_PLACEHOLDER_CSL = """<?xml version="1.0" encoding="utf-8"?>
<!-- TODO: 由 thesisforge template migrate 生成的占位 CSL，请替换为 v0.3
     citation.style 对应的真实 CSL 1.0.1 文件，并在 provenance.yaml 记录其
     哈希与许可（SCHEMA §1.1/§3.19）。 -->
<style xmlns="http://purl.org/net/xbiblio/csl" class="in-text" version="1.0.1"
       default-locale="zh-CN">
  <info>
    <title>TODO Placeholder Citation Style</title>
    <id>http://www.zotero.org/styles/thesisforge-migrate-placeholder</id>
    <updated>2000-01-01T00:00:00+00:00</updated>
  </info>
  <citation>
    <layout delimiter="; ">
      <text variable="title"/>
    </layout>
  </citation>
  <bibliography>
    <layout>
      <text variable="title"/>
    </layout>
  </bibliography>
</style>
"""


def _write_skeleton_files(
    template: ThesisTemplate,
    output_dir: Path,
    ledger: list[LedgerEntry],
    *,
    today: date,
) -> None:
    (output_dir / "provenance.yaml").write_text(
        _provenance_yaml(template, today=today), encoding="utf-8"
    )
    manual_items = "\n".join(
        f"- `{entry.field}`：{entry.reason}"
        + (f"（建议：{entry.suggestion}）" if entry.suggestion else "")
        for entry in ledger
        if entry.status == MANUAL_REQUIRED
    )
    package_id, _ = _sanitize_id(template.id)
    (output_dir / "README.md").write_text(
        _README_TEMPLATE.format(
            name=template.name,
            legacy_id=template.id,
            package_id=package_id,
            version=GENERATED_VERSION,
            manual_items=manual_items or "- 无",
        ),
        encoding="utf-8",
    )
    (output_dir / "CHANGELOG.md").write_text(
        f"# Changelog\n\n## {GENERATED_VERSION}\n\n"
        f"- 由 v0.3 模板 `{template.id}` 迁移生成（schema_version 2）。\n",
        encoding="utf-8",
    )
    # 偏差说明：§8.2 原文「fixtures 不生成」，但 fixtures/minimal 为 §1.1 必需
    # （L1 missing-package-file error），故生成最小骨架并在台账标记人工补齐。
    minimal = output_dir / "fixtures" / "minimal"
    minimal.mkdir(parents=True, exist_ok=True)
    (minimal / "thesis.md").write_text(_FIXTURE_MD, encoding="utf-8")


# ---------------------------------------------------------------------------
# 编排
# ---------------------------------------------------------------------------


def migrate_template(
    source: str | Path,
    output_dir: str | Path,
    *,
    force: bool = False,
    today: date | None = None,
) -> MigrateReport:
    """执行迁移；前置失败抛 `MigrateError` / `TemplateLoadError`。"""
    source_path = Path(source).expanduser()
    if not source_path.is_absolute():
        source_path = Path.cwd() / source_path
    output = Path(output_dir).expanduser()
    if not output.is_absolute():
        output = Path.cwd() / output

    if (
        output.exists()
        and (not output.is_dir() or any(output.iterdir()))
        and not force
    ):
        raise MigrateError(f"输出目录非空，拒绝覆盖（--force 显式开启）：{output}")

    template = load_template(source_path)  # v0.3 解析失败抛 TemplateLoadError
    raw = yaml.safe_load(source_path.read_text(encoding="utf-8"))
    if not isinstance(raw, dict):
        raw = {}

    ledger: list[LedgerEntry] = []
    template_data = _build_template_data(template, raw, ledger)

    output.mkdir(parents=True, exist_ok=True)
    header = (
        "# 由 thesisforge template migrate 生成（v0.3 → v2，SCHEMA §8）；\n"
        "# 台账见 migration-report.json，manual-required 项需人工核对。\n"
    )
    (output / "template.yaml").write_text(
        header + yaml.safe_dump(template_data, allow_unicode=True, sort_keys=False),
        encoding="utf-8",
    )
    reference_styles = _build_reference_docx(template, output / "reference.docx")
    _write_skeleton_files(
        template, output, ledger, today=today or datetime.now(UTC).date()
    )
    if "bibliography" in template_data:
        citations_dir = output / "citations"
        citations_dir.mkdir(parents=True, exist_ok=True)
        (citations_dir / "style.csl").write_text(_PLACEHOLDER_CSL, encoding="utf-8")

    # §8.2 第 4 条：产物立即跑 L1–L3
    lint_report = _lint.lint_package(output)
    report = MigrateReport(
        source=source_path,
        output=output,
        entries=tuple(ledger),
        lint_report=lint_report,
        reference_styles=tuple(reference_styles),
    )
    (output / "migration-report.json").write_text(
        json.dumps(report.to_dict(), ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    return report
