#!/usr/bin/env python3
"""Phase 0 spike：构造带锚点的 shell.docx（Template Package v2 路线②）。

以 package-sample/reference.docx 为底（继承 TF* 样式 / A4 页面设置 / 默认页眉页脚），
搭建学校前置页 + 双 section 骨架：

- section 1（front，lowerRoman 从 1 起）：封面（校名占位、logo 占位图、
  题目/姓名/学号元数据表格）、原创性声明页、目录区（tf_toc 书签锚点）；
- section 2（main，decimal 从 1 重启，独立页眉页脚）：正文区仅一个空段落，
  其上挂唯一书签 tf_body 作为编译产物插入锚点。

产出：
- package-sample/shell.docx
- package-sample/assets/logo.png  占位图（纯 stdlib 生成，确定性字节）
- output/shell-summary.json  结构化实证结论

运行（仓库根目录）：.venv/bin/python spikes/phase0/docx-template/build_shell.py
"""

from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

sys.path.insert(0, str(Path(__file__).resolve().parent))
from spike_common import (
    ASSETS_DIR,
    OUTPUT_DIR,
    PACKAGE_DIR,
    ensure_dirs,
    run_openxml_validate,
    soffice_smoke,
    summarize_validation,
    write_placeholder_png,
)

SHELL_PATH = PACKAGE_DIR / "shell.docx"
REFERENCE_PATH = PACKAGE_DIR / "reference.docx"
LOGO_PATH = ASSETS_DIR / "logo.png"
SUMMARY_PATH = OUTPUT_DIR / "shell-summary.json"

BODY_ANCHOR = "tf_body"
TOC_ANCHOR = "tf_toc"

# w:sectPr 子元素的 schema 顺序中排在 pgNumType 之后的元素，
# 插入 pgNumType 时必须放在这些元素之前，否则 Word 打开可能提示修复。
SECTPR_PGNUM_SUCCESSORS = (
    "w:cols",
    "w:formProt",
    "w:vAlign",
    "w:noEndnote",
    "w:titlePg",
    "w:textDirection",
    "w:bidi",
    "w:rtlGutter",
    "w:docGrid",
    "w:printerSettings",
    "w:sectPrChange",
)


def _set_page_number_format(section, *, fmt: str, start: int) -> None:
    """给 section 的 sectPr 写入 w:pgNumType（页码格式 + 起始值）。"""
    sect_pr = section._sectPr
    existing = sect_pr.find(qn("w:pgNumType"))
    if existing is not None:
        sect_pr.remove(existing)
    pg_num = OxmlElement("w:pgNumType")
    pg_num.set(qn("w:fmt"), fmt)
    pg_num.set(qn("w:start"), str(start))
    sect_pr.insert_element_before(pg_num, *SECTPR_PGNUM_SUCCESSORS)


def _add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    """在段落内容外侧包裹书签（bookmarkStart 须位于 pPr 之后）。"""
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    p_element = paragraph._p
    insert_at = 1 if p_element.find(qn("w:pPr")) is not None else 0
    p_element.insert(insert_at, start)
    p_element.append(end)


def _fill_header_with_border(section, text: str) -> None:
    """页眉：居中校名文本 + 0.5pt 下边框线（与 reference.docx 页眉同款）。"""
    paragraph = section.header.paragraphs[0]
    paragraph.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")  # 0.5pt，单位 1/8pt
    bottom.set(qn("w:color"), "auto")
    bottom.set(qn("w:space"), "1")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _fill_footer_with_page_field(section) -> None:
    """页脚：居中 PAGE 域（页码格式由所在 section 的 pgNumType 决定）。"""
    paragraph = section.footer.paragraphs[0]
    paragraph.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    begin = OxmlElement("w:r")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    begin.append(fld_begin)
    instr = OxmlElement("w:r")
    instr_text = OxmlElement("w:instrText")
    instr_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr_text.text = " PAGE "
    instr.append(instr_text)
    end = OxmlElement("w:r")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    end.append(fld_end)
    paragraph._p.extend((begin, instr, end))


def _add_cover(document, logo_path: Path) -> None:
    """封面页：校名占位、logo 占位图、论文元数据表格占位。"""
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("湖南工业大学")
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    document.add_picture(str(logo_path), width=Mm(35))
    logo_paragraph = document.paragraphs[-1]
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    degree = document.add_paragraph()
    degree.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = degree.add_run("硕士学位论文")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    table = document.add_table(rows=4, cols=2)
    table.style = document.styles["Table Grid"]
    placeholders = (
        ("论文题目", "【论文题目占位】"),
        ("学生姓名", "【姓名占位】"),
        ("学号", "【学号占位】"),
        ("指导教师", "【导师占位】"),
    )
    for row, (label, value) in zip(table.rows, placeholders, strict=True):
        row.cells[0].text = label
        row.cells[1].text = value


def _add_declaration(document) -> None:
    """原创性声明页（占位文本）。"""
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("原创性声明")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    document.add_paragraph(
        "本人郑重声明：所呈交的学位论文，是本人在导师指导下独立进行研究工作"
        "所取得的成果。除文中已经注明引用的内容外，本论文不包含任何其他个人"
        "或集体已经发表或撰写过的研究成果。（占位文本，正式内容以学校文件为准）"
    )
    document.add_paragraph("论文作者签名：______________    日期：____年____月____日")


def build_shell() -> dict:
    """生成 shell.docx，返回结构摘要。"""
    if not REFERENCE_PATH.is_file():
        raise SystemExit(f"缺少 {REFERENCE_PATH}，请先运行 build_reference.py")
    write_placeholder_png(LOGO_PATH, width=480, height=240)

    document = Document(REFERENCE_PATH)  # 继承 reference.docx 的样式/页面/页眉页脚

    _add_cover(document, LOGO_PATH)
    document.add_page_break()
    _add_declaration(document)
    document.add_page_break()

    toc_title = document.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_title.add_run("目  录")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    toc_anchor = document.add_paragraph()
    _add_bookmark(toc_anchor, TOC_ANCHOR, 1001)

    # 分节：add_section 会把当前 sectPr 克隆进上一节末尾的分节符段落，
    # 原 final sectPr 成为新 main section。
    main_section = document.add_section(WD_SECTION.NEW_PAGE)
    body_anchor = document.add_paragraph()
    _add_bookmark(body_anchor, BODY_ANCHOR, 1002)

    front_section = document.sections[0]
    _set_page_number_format(front_section, fmt="lowerRoman", start=1)
    _set_page_number_format(main_section, fmt="decimal", start=1)

    # main section 独立页眉页脚；front section 沿用 reference.docx 的页眉页脚
    main_section.header.is_linked_to_previous = False
    main_section.footer.is_linked_to_previous = False
    _fill_header_with_border(main_section, "湖南工业大学硕士学位论文")
    _fill_footer_with_page_field(main_section)

    document.save(SHELL_PATH)
    return {
        "sections": len(document.sections),
        "anchors": [TOC_ANCHOR, BODY_ANCHOR],
        "logo": str(LOGO_PATH),
    }


def verify_shell() -> dict:
    """结构断言：双 section、页码格式、锚点唯一且配对、图片 relationship。"""
    evidence: dict[str, object] = {}

    document = Document(SHELL_PATH)
    assert len(document.sections) == 2

    section_info = []
    for section in document.sections:
        sect_pr = section._sectPr
        pg_num = sect_pr.find(qn("w:pgNumType"))
        refs = [
            (el.tag.rsplit("}", 1)[-1], el.get(qn("w:type")), el.get(qn("r:id")))
            for el in sect_pr
            if el.tag.rsplit("}", 1)[-1] in ("headerReference", "footerReference")
        ]
        section_info.append(
            {
                "pg_num_fmt": None if pg_num is None else pg_num.get(qn("w:fmt")),
                "pg_num_start": None if pg_num is None else pg_num.get(qn("w:start")),
                "header_footer_refs": refs,
            }
        )
    evidence["sections"] = section_info
    assert section_info[0]["pg_num_fmt"] == "lowerRoman"
    assert section_info[0]["pg_num_start"] == "1"
    assert section_info[1]["pg_num_fmt"] == "decimal"
    assert section_info[1]["pg_num_start"] == "1"
    # 两节都必须有自己的页眉页脚引用（front 继承 reference，main 独立）
    assert section_info[0]["header_footer_refs"]
    assert section_info[1]["header_footer_refs"]

    body = document.element.body
    for anchor in (TOC_ANCHOR, BODY_ANCHOR):
        starts = [el for el in body.iter(qn("w:bookmarkStart")) if el.get(qn("w:name")) == anchor]
        assert len(starts) == 1, f"锚点 {anchor} 应唯一，实际 {len(starts)} 个"
    evidence["anchors_unique"] = True

    rels = document.part.rels
    image_rels = [r for r in rels.values() if "image" in r.reltype]
    assert len(image_rels) == 1
    evidence["image_relationship"] = {
        "rId": image_rels[0].rId,
        "target": image_rels[0].target_ref,
    }

    header_text = "".join(
        node.text or "" for node in document.sections[1].header.part.element.iter(qn("w:t"))
    )
    evidence["main_header_text"] = header_text
    assert "湖南工业大学硕士学位论文" in header_text
    return evidence


def main() -> None:
    ensure_dirs()
    built = build_shell()
    evidence = verify_shell()
    validation = run_openxml_validate(SHELL_PATH)
    smoke = soffice_smoke(SHELL_PATH, work_dir=OUTPUT_DIR / "shell-smoke")

    summary = {
        "shell_docx": str(SHELL_PATH),
        "built": built,
        "validation": summarize_validation(validation),
        "soffice_smoke": smoke,
        "evidence": evidence,
    }
    SUMMARY_PATH.write_text(
        json.dumps(summary, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    print(json.dumps(summary, ensure_ascii=False, indent=2))
    if not validation["ok"]:
        raise SystemExit("openxml 校验未通过，见上方 summary")
    if smoke.get("available") and not smoke["ok"]:
        raise SystemExit("soffice 冒烟转换失败，见上方 summary")


if __name__ == "__main__":
    main()
