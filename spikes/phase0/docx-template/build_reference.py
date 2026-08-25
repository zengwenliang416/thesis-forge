#!/usr/bin/env python3
"""Phase 0 spike：以 HUT 学校 YAML 为蓝本编程生成 reference.docx。

验证 Template Package v2 路线①：用 reference.docx 作为样式/主题/页面设置来源，
替代 python-docx 默认模板。

产出：
- package-sample/reference.docx  空正文，含 TF* 样式、页面设置、页眉页脚
- output/reference-inheritance.docx  以 reference.docx 为底新建文档的继承验证
- output/reference-summary.json  结构化实证结论

运行（仓库根目录）：.venv/bin/python spikes/phase0/docx-template/build_reference.py
"""

from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

sys.path.insert(0, str(Path(__file__).resolve().parent))
from spike_common import (
    OUTPUT_DIR,
    PACKAGE_DIR,
    REPO_ROOT,
    ensure_dirs,
    run_openxml_validate,
    summarize_validation,
)

from thesis_forge.renderers.docx.document import configure_section_geometry
from thesis_forge.renderers.docx.styles import (
    apply_paragraph_style,
    configure_styles,
)
from thesis_forge.templates import load_template
from thesis_forge.templates.model import ParagraphStyleSpec

HUT_YAML = (
    REPO_ROOT / "templates" / "schools" / "hunan-university-of-technology" / "master-2026.yaml"
)

REFERENCE_PATH = PACKAGE_DIR / "reference.docx"
INHERITANCE_PATH = OUTPUT_DIR / "reference-inheritance.docx"
SUMMARY_PATH = OUTPUT_DIR / "reference-summary.json"

TWIPS_PER_PT = 20


def _add_paragraph_style(document, name: str, base=None):
    style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = base if base is not None else document.styles["Normal"]
    style.quick_style = True
    return style


def _set_default_fonts(document, *, east_asia: str, latin: str, size_pt: float) -> None:
    """把 docDefaults 的字体/字号改成学校基线，替代 python-docx 默认主题字体。"""
    styles_element = document.styles.element
    doc_defaults = styles_element.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles_element.insert(0, doc_defaults)
    r_pr_default = doc_defaults.find(qn("w:rPrDefault"))
    if r_pr_default is None:
        r_pr_default = OxmlElement("w:rPrDefault")
        doc_defaults.append(r_pr_default)
    r_pr = r_pr_default.find(qn("w:rPr"))
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        r_pr_default.append(r_pr)
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        r_fonts.attrib.pop(qn(f"w:{attr}"), None)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:cs"), latin)
    sz = r_pr.find(qn("w:sz"))
    if sz is None:
        sz = OxmlElement("w:sz")
        r_pr.append(sz)
    sz.set(qn("w:val"), str(int(size_pt * 2)))


def _declare_font_table(document, font_names: list[str]) -> None:
    """在 fontTable 中登记学校字体（python-docx 无公开 API，直接改部件 XML）。

    python-docx 1.2.0 把 fontTable.xml 加载为普通 Part（无 _element），
    必须经 blob 反序列化/回写。这是 reference.docx 路线的实证坑之一。
    """
    part = next(
        part for part in document.part.package.parts if str(part.partname) == "/word/fontTable.xml"
    )
    root = etree.fromstring(part.blob)
    existing = {el.get(qn("w:name")) for el in root.findall(qn("w:font"))}
    for name in font_names:
        if name in existing:
            continue
        font = OxmlElement("w:font")
        font.set(qn("w:name"), name)
        family = OxmlElement("w:family")
        family.set(qn("w:val"), "roman")
        font.append(family)
        charset = OxmlElement("w:charset")
        charset.set(qn("w:val"), "86")  # GB2312，中文字体常规声明
        font.append(charset)
        root.append(font)
    part._blob = etree.tostring(
        root,
        xml_declaration=True,
        encoding="UTF-8",
        standalone=True,
    )


def _build_header_footer(document, header_text: str) -> None:
    """reference.docx 自带默认页眉（校名+下边框线）与页脚（PAGE 域居中）。"""
    section = document.sections[0]
    header_paragraph = section.header.paragraphs[0]
    header_paragraph.style = document.styles["Header"]
    run = header_paragraph.add_run(header_text)
    run.font.size = None  # 继承 Header 样式
    p_pr = header_paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")  # 0.5pt，单位 1/8pt
    bottom.set(qn("w:color"), "auto")
    bottom.set(qn("w:space"), "1")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    footer_paragraph = section.footer.paragraphs[0]
    footer_paragraph.style = document.styles["Footer"]
    begin = OxmlElement("w:r")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    begin.append(fld_begin)
    instr = OxmlElement("w:r")
    instr_text = OxmlElement("w:instrText")
    instr_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr_text.text = " PAGE "
    instr.append(instr_text)
    end = OxmlElement("w:r")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    end.append(fld_end)
    footer_paragraph._p.extend((begin, instr, end))


def build_reference() -> dict:
    """生成 reference.docx，返回样式清单。"""
    template = load_template(HUT_YAML)
    document = Document()  # 以 python-docx 默认包为骨架，注入学校样式
    configure_section_geometry(document.sections[0], template.page)

    _set_default_fonts(
        document,
        east_asia=template.body.font.east_asia,
        latin=template.body.font.latin,
        size_pt=12.0,
    )

    normal = document.styles["Normal"]
    base_font_only = ParagraphStyleSpec(
        font=template.body.font,
        size=template.body.size,
    )
    apply_paragraph_style(normal, base_font_only)

    body = _add_paragraph_style(document, "TF Body")
    apply_paragraph_style(body, template.body)

    body_first = _add_paragraph_style(document, "TF Body First", base=body)
    apply_paragraph_style(
        body_first,
        ParagraphStyleSpec(first_line_indent=template.heading.level1.first_line_indent),
    )

    heading_specs = {
        1: template.heading.level1,
        2: template.heading.level2,
        3: template.heading.level3,
        4: template.heading.level3,  # HUT YAML 只定义到三级，四级样例复用三级
    }
    for level, spec in heading_specs.items():
        style = _add_paragraph_style(document, f"TF Heading {level}")
        apply_paragraph_style(
            style,
            spec,
            fallback_font=template.body.font,
            fallback_size=template.body.size,
        )

    abstract_body = (
        template.semantic_styles.abstract_zh.body
        if template.semantic_styles.abstract_zh is not None
        and template.semantic_styles.abstract_zh.body is not None
        else template.body
    )
    abstract = _add_paragraph_style(document, "TF Abstract")
    apply_paragraph_style(abstract, abstract_body)

    bibliography = _add_paragraph_style(document, "TF Bibliography")
    apply_paragraph_style(
        bibliography,
        template.bibliography.entry,
        fallback_font=template.body.font,
        fallback_size=template.body.size,
    )

    figure_caption = _add_paragraph_style(document, "TF Figure Caption")
    caption = template.figure.caption
    apply_paragraph_style(
        figure_caption,
        ParagraphStyleSpec(font=caption.font, size=caption.size, alignment=caption.alignment),
    )

    table_caption = _add_paragraph_style(document, "TF Table Caption")
    caption = template.table.caption
    apply_paragraph_style(
        table_caption,
        ParagraphStyleSpec(font=caption.font, size=caption.size, alignment=caption.alignment),
    )

    equation = _add_paragraph_style(document, "TF Equation")
    apply_paragraph_style(
        equation,
        ParagraphStyleSpec(
            font=template.body.font,
            size=template.body.size,
            alignment="center",
            first_line_indent=template.heading.level1.first_line_indent,
            line_spacing=template.body.line_spacing,
        ),
    )

    code_char = document.styles.add_style("TF Code Char", WD_STYLE_TYPE.CHARACTER)
    code_char.font.name = "Consolas"

    _declare_font_table(
        document,
        [template.body.font.east_asia, "黑体", template.body.font.latin],
    )
    _build_header_footer(document, "湖南工业大学硕士学位论文")

    document.save(REFERENCE_PATH)
    return {
        "styles": sorted(
            style.name for style in document.styles if style.name and style.name.startswith("TF ")
        )
    }


def verify_inheritance() -> dict:
    """以 reference.docx 为底新建文档，验证样式/页面/页眉页脚可继承。"""
    evidence: dict[str, object] = {}

    default_doc = Document()
    default_names = {style.name for style in default_doc.styles}
    evidence["default_template_has_tf_styles"] = "TF Body" in default_names

    document = Document(REFERENCE_PATH)
    style_names = {style.name for style in document.styles}
    expected = {
        "TF Body",
        "TF Body First",
        "TF Heading 1",
        "TF Heading 2",
        "TF Heading 3",
        "TF Heading 4",
        "TF Abstract",
        "TF Bibliography",
        "TF Figure Caption",
        "TF Table Caption",
        "TF Equation",
        "TF Code Char",
    }
    evidence["missing_styles"] = sorted(expected - style_names)
    assert not evidence["missing_styles"], evidence["missing_styles"]

    section = document.sections[0]
    evidence["section"] = {
        "page_width_mm": round(section.page_width.mm, 2),
        "left_margin_mm": round(section.left_margin.mm, 2),
        "header_distance_mm": round(section.header_distance.mm, 2),
        "footer_distance_mm": round(section.footer_distance.mm, 2),
        "doc_grid": [
            (el.get(qn("w:type")), el.get(qn("w:linePitch")))
            for el in section._sectPr.findall(qn("w:docGrid"))
        ],
    }
    assert round(section.left_margin.mm) == 30
    assert round(section.header_distance.mm) == 15

    header_text = "".join(node.text or "" for node in section.header.part.element.iter(qn("w:t")))
    evidence["header_text"] = header_text
    assert "湖南工业大学" in header_text

    # 用继承来的样式排版新内容
    document.add_paragraph("第一章 绪论", style="TF Heading 1")
    document.add_paragraph("正文段落，验证 TF Body 的继承效果。", style="TF Body")
    document.add_paragraph("图 1-1  示例", style="TF Figure Caption")
    document.add_paragraph("[1] 参考文献条目。", style="TF Bibliography")

    document.save(INHERITANCE_PATH)

    # 关键限制实证：现有渲染器 configure_styles 会把 YAML 值直接写进样式。
    # 即 reference.docx 里的同名样式会被渲染器改写而非保留 —— 优先级必须显式定义。
    template = load_template(HUT_YAML)
    probe = Document(REFERENCE_PATH)
    normal_before = probe.styles["Normal"].paragraph_format.first_line_indent
    configure_styles(probe, template)
    normal_after = probe.styles["Normal"].paragraph_format.first_line_indent
    evidence["renderer_overwrites_styles"] = {
        "normal_first_line_indent_before": (None if normal_before is None else normal_before.twips),
        "normal_first_line_indent_after": (None if normal_after is None else normal_after.twips),
        "conclusion": "configure_styles 直接改写 Normal/Heading/TOC 样式属性，"
        "reference.docx 与 YAML 并存时必须定义优先级",
    }
    assert normal_after is not None and normal_after.twips == 480

    # python-docx 以 docx 起建 = 就地编辑该文件，保存需另存路径
    evidence["opens_in_place"] = (
        "Document(path) 打开的是文件本体，document.save() 不带参数会覆盖源文件；"
        "作为模板使用必须先复制或另存"
    )
    return evidence


def main() -> None:
    ensure_dirs()
    built = build_reference()

    reference_report = run_openxml_validate(REFERENCE_PATH)
    inheritance_report = run_openxml_validate(INHERITANCE_PATH) if False else None
    evidence = verify_inheritance()
    inheritance_report = run_openxml_validate(INHERITANCE_PATH)

    summary = {
        "reference_docx": str(REFERENCE_PATH),
        "reference_validation": summarize_validation(reference_report),
        "inheritance_docx": str(INHERITANCE_PATH),
        "inheritance_validation": summarize_validation(inheritance_report),
        "tf_styles": built["styles"],
        "evidence": evidence,
    }
    SUMMARY_PATH.write_text(
        json.dumps(summary, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    print(json.dumps(summary, ensure_ascii=False, indent=2))
    ok = reference_report["ok"] and inheritance_report["ok"]
    if not ok:
        raise SystemExit("openxml 校验未通过，见上方 summary")


if __name__ == "__main__":
    main()
